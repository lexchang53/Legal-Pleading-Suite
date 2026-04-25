import sys
import os
import argparse
from pathlib import Path

# 初始化：強制標準輸出採用 UTF-8，避免 Windows (CP950) 環境下印出中文字元出錯
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Cm, Twips
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

import header_utils

from lxml import etree
from markdown_parser import Block, TableBlock, auto_fix_markdown, parse_markdown
from word_xml_utils import (
    find_and_remove_anchor, clear_body, enable_line_numbering,
    disable_hanging_punctuation, create_override_num, create_l2_reset_num, set_num_pr
)

# ==============================================================================
# 環境與依賴準備
# ==============================================================================

# 設定對應的技能路徑
_script_dir = os.path.dirname(os.path.abspath(__file__))

try:
    sys.path.insert(0, _script_dir)
    import check_citations
    _HAS_CITATION_CHECK = True
except ImportError:
    _HAS_CITATION_CHECK = False

# ==============================================================================
# 常數定義
# ==============================================================================

BODY_INDENT_MAP = {
    0: 567,   # 第一層（一、）：往內縮兩個中文字（約0.8cm~1cm）
    1: 850,   # 第二層 (一) ：往內縮三個中文字
    2: 850,   # 第三層 1.
    3: 1134,  # 第四層 (1)
    4: 1134,  # 更深層
}

# 狀首專用樣式
HEADER_STYLES = {'書狀_標題', '書狀_狀首日期', '書狀_預設'}

# 目標樣式清單
TARGET_STYLES = [
    '書狀_預設', '書狀_標題', '書狀_簽章',
    '書狀_狀首日期', '書狀_狀尾日期', '書狀_謹狀',
    '書狀_被上證據編號', '書狀_被上證據編號10', '書狀_證據編號', '書狀_證據編號10',
    '書狀_狀首當事人', '書狀_清單',
    '通用_層級1', '通用_層級2', '通用_層級3', '通用_層級4'
]

# ==============================================================================
# 核心函式
# ==============================================================================

def has_bold_marker(text):
    import re
    return bool(re.search(r'\*\*|__', text))

def split_bold_runs(text):
    """
    將帶有 **粗體** 標記的字串切分為 (text, is_bold) tuple 列表。
    """
    import re
    parts = re.split(r'(\*\*|__)', text)
    runs = []
    in_bold = False
    for p in parts:
        if p in ('**', '__'):
            in_bold = not in_bold
        else:
            if p:
                runs.append((p, in_bold))
    return runs

def _add_paragraph_with_bold(doc, text, style):
    p = doc.add_paragraph(style=style)
    if not has_bold_marker(text):
        p.add_run(text)
    else:
        for run_text, is_bold in split_bold_runs(text):
            if run_text:
                run = p.add_run(run_text)
                if is_bold:
                    run.bold = True
    return p

def write_paragraph(doc, block, num_id, outline_level=None):
    p = _add_paragraph_with_bold(doc, block.text, block.style)

    if block.needs_num and block.ilvl is not None:
        set_num_pr(p, num_id, block.ilvl, outline_level=outline_level)

    return p

# ==============================================================================
# 模板檢查
# ==============================================================================

def check_template(template_path):
    print(f"正在檢查模板: {template_path}\n")
    doc = Document(template_path)

    existing_styles = {s.name for s in doc.styles}
    all_ok = True

    print("=== 樣式完整性 ===")
    for ts in TARGET_STYLES:
        if ts in existing_styles:
            print(f"  OK: {ts}")
        else:
            print(f"  MISSING: {ts}")
            all_ok = False

    print("\n=== 錨點段落 ===")
    try:
        num_id, abstract_num_id = find_and_remove_anchor(doc)
        print(f"  OK: numId={num_id}, abstractNumId={abstract_num_id}")
    except RuntimeError as e:
        print(f"  ERROR: {e}")
        all_ok = False

    print(f"\n{'=== 檢查通過 ===' if all_ok else '=== 檢查失敗 ==='}")
    return all_ok


def write_table(doc, table_block):
    table = doc.add_table(rows=1 + len(table_block.rows), cols=3)
    table.autofit = False
    table.allow_autofit = False
    
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr[0].append(tblLayout)

    widths = [Cm(2.3), Cm(11.5), Cm(2.2)]
    for i, col in enumerate(table.columns):
        col.width = widths[i]
        
    for i, row in enumerate(table.rows):
        is_header = (i == 0)
        trPr = row._tr.get_or_add_trPr()
        
        if is_header:
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
            row.allow_break_across_pages = False
        
        cells_data = table_block.headers if is_header else table_block.rows[i-1]
        for j, cell in enumerate(row.cells):
            if j < len(widths):
                cell.width = widths[j]
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4') 
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'auto')
            tcPr.append(shd)
            vAlign = OxmlElement('w:vAlign')
            if is_header or j == 0 or j == 2:
                vAlign.set(qn('w:val'), 'center')
            else:
                vAlign.set(qn('w:val'), 'top')
            tcPr.append(vAlign)
            
            text = cells_data[j] if j < len(cells_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = '書狀_預設'
            if text:
                run = p.add_run(text)
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '標楷體')
                rFonts.set(qn('w:ascii'), '標楷體')
                rFonts.set(qn('w:hAnsi'), '標楷體')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '28')
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '28')
                rPr.append(szCs)
            if is_header or j == 0 or j == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pPr_cell = p._element.get_or_add_pPr()
            spacing_el = OxmlElement('w:spacing')
            spacing_el.set(qn('w:line'), '240')       
            spacing_el.set(qn('w:lineRule'), 'auto')  
            pPr_cell.append(spacing_el)

    all_rows = table.rows
    for row_idx, row in enumerate(all_rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                if row_idx < len(all_rows) - 1:
                    para.paragraph_format.keep_with_next = True

# ==============================================================================
# 主程式
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(description='draft-pleading 書狀排版工具')
    parser.add_argument('draft', nargs='?', help='Markdown 草稿路徑')
    parser.add_argument('--template', help='模板路徑',
                        default=os.path.join(os.path.dirname(__file__),
                                             '..', 'assets', 'pleading-tmpl.docx'))
    parser.add_argument('--header', help='狀首來源 .docx 路徑')
    parser.add_argument('--output', help='輸出路徑')
    parser.add_argument('--check-template', dest='check', metavar='PATH',
                        help='僅檢查模板完整性')
    args = parser.parse_args()

    if args.check:
        sys.exit(0 if check_template(args.check) else 1)

    if not args.draft:
        parser.error("請提供 Markdown 草稿路徑")

    template_path = os.path.abspath(args.template)
    draft_path = os.path.abspath(args.draft)

    if not os.path.exists(template_path):
        print(f"[ERROR] 模板不存在: {template_path}")
        sys.exit(1)
    if not os.path.exists(draft_path):
        print(f"[ERROR] 草稿不存在: {draft_path}")
        sys.exit(1)

    if args.output:
        output_path = os.path.abspath(args.output)
    else:
        base = os.path.splitext(os.path.basename(draft_path))[0]
        output_dir = os.path.dirname(draft_path)
        output_path = os.path.join(output_dir, f"{base}_排版完成.docx")

    print(f"[START] 開始排版 (v2 iteration-3)")
    print(f"  模板: {template_path}")
    print(f"  草稿: {draft_path}")
    print(f"  輸出: {output_path}")
    if args.header:
        print(f"  狀首: {args.header}")
    print()

    if _HAS_CITATION_CHECK:
        print("[CHECK] 執行引用規則預檢...")
        ok = check_citations.run_check(draft_path, strict=True)
        if not ok:
            print()
            print("[ABORT] 草稿存在引用規則違規，排版已中止。")
            print("        請依照上方報告修正草稿後，再次執行排版。")
            sys.exit(1)
        print("[CHECK] 引用規則預檢通過。")
        print()
    else:
        print("[WARN] check_citations 模組未載入，跳過引用規則預檢。")
        print()

    doc = Document(template_path)

    anchor_num_id, abstract_num_id = find_and_remove_anchor(doc)

    clear_body(doc)

    fixed_content = auto_fix_markdown(draft_path)

    blocks = parse_markdown(draft_path, content=fixed_content)

    md_headers = header_utils.extract_md_header_info(blocks)
    if args.header:
        header_path = os.path.abspath(args.header)
        if os.path.exists(header_path):
            header_data = header_utils.extract_header(header_path)
            header_utils.merge_and_write_header(doc, header_data, md_headers, is_issue_table=False)
        else:
            print(f"[WARN] 狀首來源不存在: {header_path}")
            header_utils.merge_and_write_header(doc, None, md_headers, is_issue_table=False)
    else:
        header_utils.merge_and_write_header(doc, None, md_headers, is_issue_table=False)

    current_num_id = anchor_num_id
    para_count = 0
    last_heading_ilvl = None  

    import re
    has_declaration = any(
        getattr(b, 'is_semantic_heading', False) and b.text.strip().replace('*', '').endswith('聲明')
        for b in blocks
    )
    current_section = '聲明' if has_declaration else None

    skip_header_blocks = True

    in_signature_block = False

    _need_l2_reset = False
    _last_l1_num_id = current_num_id  

    for block in blocks:
        if getattr(block, 'style', '') == 'TABLE' or isinstance(block, TableBlock):
            write_table(doc, block)
            para_count += 1
            last_heading_ilvl = None
            continue

        if skip_header_blocks:
            if '為就' in block.text and re.search(r'事[：:]\s*$', block.text):
                skip_header_blocks = False
            continue

        if block.style == '書狀_謹狀':
            in_signature_block = True
            blank = doc.add_paragraph('')
            blank.style = doc.styles['書狀_預設']
            para_count += 1

        if block.is_override_trigger:
            current_num_id = create_override_num(doc, abstract_num_id)
            _last_l1_num_id = current_num_id
            _need_l2_reset = False

        if getattr(block, 'is_semantic_heading', False):
            p = _add_paragraph_with_bold(doc, block.text, block.style)
            
            text_clean = block.text.replace('*', '').replace('\u3000', '').replace(' ', '').strip()
            if text_clean.endswith('聲明') or text_clean == '聲明事項':
                current_section = '聲明'
            elif text_clean in ('事實與理由', '理由'):
                current_section = '理由'

            if current_section == '理由':
                p.paragraph_format.keep_with_next = True
            else:
                p.paragraph_format.keep_with_next = False
                
            last_heading_ilvl = None  

        elif block.needs_num and block.ilvl is not None:
            if block.ilvl == 0:
                _need_l2_reset = True
                _last_l1_num_id = current_num_id
                effective_num_id = current_num_id
            elif block.ilvl == 1 and _need_l2_reset:
                current_num_id = create_l2_reset_num(doc, abstract_num_id, _last_l1_num_id)
                _need_l2_reset = False  
                effective_num_id = current_num_id
            else:
                effective_num_id = current_num_id
            
            effective_outline_level = None
            if block.ilvl == 0 and current_section != '聲明':
                if current_section == '理由' or current_section is None:
                    effective_outline_level = 0

            p = write_paragraph(doc, block, effective_num_id, outline_level=effective_outline_level)
            
            if block.ilvl == 0:
                if current_section == '聲明':
                    p.paragraph_format.keep_with_next = False
                else:
                    p.paragraph_format.keep_with_next = True
            
            last_heading_ilvl = block.ilvl  

            should_bold = False
            if getattr(block, 'has_child', False) and current_section != '聲明':
                should_bold = True
            
            if should_bold:
                for run in p.runs:
                    run.bold = True

        elif block.style == '書狀_預設' and last_heading_ilvl is not None:
            if block.text.strip() == '證據或文件清單：':
                blank = doc.add_paragraph()
                blank.paragraph_format.keep_with_next = True
                p = _add_paragraph_with_bold(doc, block.text, block.style)
                p.paragraph_format.keep_with_next = True
                last_heading_ilvl = None
                p.paragraph_format.left_indent = Twips(0)
            else:
                p = _add_paragraph_with_bold(doc, block.text, block.style)
                p.paragraph_format.left_indent = Twips(BODY_INDENT_MAP[last_heading_ilvl])
        else:
            if block.text.strip() == '證據或文件清單：':
                blank = doc.add_paragraph()
                blank.paragraph_format.keep_with_next = True
                p = _add_paragraph_with_bold(doc, block.text, block.style)
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.left_indent = Twips(0)
            else:
                p = _add_paragraph_with_bold(doc, block.text, block.style)
                
            last_heading_ilvl = None  

        if in_signature_block:
            pPr = p._element.get_or_add_pPr()
            if pPr.find(qn('w:keepLines')) is None:
                etree.SubElement(pPr, qn('w:keepLines'))
            if pPr.find(qn('w:keepNext')) is None:
                etree.SubElement(pPr, qn('w:keepNext'))

        para_count += 1

    enable_line_numbering(doc)
    
    disable_hanging_punctuation(doc)

    doc.save(output_path)
    print(f"\n[DONE] 排版完成！")
    print(f"  總段落數: {para_count}")
    print(f"  輸出檔案: {output_path}")

if __name__ == '__main__':
    main()
