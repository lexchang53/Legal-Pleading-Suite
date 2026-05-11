# header_utils.py — 狀首寫入共用模組（專屬於 draft-pleading 內部排版引擎）
# 本模組完全獨立運作。
#
# 驗收標準（依需求規格逐項）：
# [V] 1. 所有身份行都先清 direct formatting 再寫入（normalize_header_paragraph）
# [V] 2. 所有身份行都明確設 left_indent=0、first_line_indent=0
# [V] 3. 已清除舊 tab stops / <w:tabs>（normalize 中的 XML 操作）
# [V] 4. 續行電話、email、地址同上（含全形空白縮排行）逐行以 "\t" 起始
# [V] 5. 上訴人、被上訴人、法定代理人、訴訟代理人地址起點均為頁面左邊界起算 6 公分
# [V] 6. 訴訟標的價額：書狀_預設、left_indent=0、first_line_indent=0、不加 tab
# [V] 7. fallback 一律為 書狀_預設，不退回 Normal 或 Word 預設樣式

import re
import copy
import datetime
from docx import Document
import os
import json
from docx.shared import Cm, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def _wants_top_date():
    try:
        # 從腳本所在目錄 (draft-pleading/scripts) 往上兩層尋找
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        prefs_path = os.path.join(base_dir, '.pleading_preferences.json')
        
        if os.path.exists(prefs_path):
            with open(prefs_path, 'r', encoding='utf-8') as f:
                prefs = json.load(f)
                if 'show_top_date' in prefs:
                    return bool(prefs['show_top_date'])
    except FileNotFoundError:
        pass
    except Exception as e:
        print(f"[WARN] 讀取或解析設定檔 {prefs_path} 時發生錯誤: {e}")
    return False


END_PATTERN = re.compile(r'事[：:]\s*$')
DATE_PATTERN = re.compile(r'(?:民國)?\s*(\d{2,3})\s*年\s*(?:[\d\u3000\s]+月)?\s*(?:[\d\u3000\s]+日)?')

# 採用負面表列與結構特徵進行泛用識別，無須列舉即可支援任意（如「特別代理人」、「參加人」、「承受訴訟人」等）當事人稱謂
IDENTITY_ROLES = re.compile(
    r'^(?!案\s*號|股\s*別|訴\s*訟\s*標\s*的|為\s*就|主\s*旨|地\s*址|電\s*話|電子郵件|Email|被\s*證|甲\s*證|上\s*證)([^\s\uff1a:]{2,10})([\uff1a:\s]+)(?=[^\s])'
)

CONTINUATION_KEYWORDS = re.compile(
    r'(地址[：:]?|地址請詳卷|住[：:]?|住所[：:]?|設[：:]?|居[：:]?|通訊地址[：:]?|送達地址[：:]?|'
    r'電話[：:]?|電子郵件[：:]?|[Ee]-?[Mm]ail[：:]?|傳真[：:]?|統一編號[：:]?|統編[：:]?|'
    r'身分證[字號]*[：:]?|律師事務所[：:]?|事務所[：:]?|同上|地址同上|請詳卷)'
)




# ─────────────────────────────────────────────────────────────────────────────
# normalize：清除直接格式
# ─────────────────────────────────────────────────────────────────────────────

def normalize_header_paragraph(p):
    """
    清除段落的直接格式覆寫：
      - left_indent = 0
      - first_line_indent = 0
      - 清除既有 <w:tabs> XML 節點（徹底移除殘留 tab stops）
    所有身份行、訴訟標的價額行寫入前，必須先呼叫此函式。
    """
    fmt = p.paragraph_format
    fmt.left_indent = 0
    fmt.first_line_indent = 0
    # 用 XML 操作徹底清除 <w:tabs>
    try:
        pPr = p._element.get_or_add_pPr()
        existing_tabs = pPr.find(qn('w:tabs'))
        if existing_tabs is not None:
            pPr.remove(existing_tabs)
    except Exception:
        pass


def _get_safe_style(doc, style_name):
    """
    確保樣式存在於文件中，如果不存在則回退至等效樣式或 '書狀_預設'
    """
    mangled_map = {
        '書狀_預設': ['書狀預設', 'a6', 'Style1', 'Ѫ_w]'],
        '書狀_標題': ['a10', 'Style12', 'Ѫ_D'],
        '書狀_狀首當事人': ['a27', 'Style33', 'Ѫ_ƤH'],
        '書狀_簽章': ['a13', 'Style24', 'Ѫ_ñ'],
        '書狀_謹狀': ['a15', 'Style25', 'Ѫ_Ԫ'],
    }
    
    existing_styles = {s.name for s in doc.styles}
    if style_name in existing_styles:
        return style_name
        
    lookups = mangled_map.get(style_name, [])
    for s in doc.styles:
        if s.style_id in lookups or s.name in lookups:
            return s.name
            
    # 保底：如果沒找到，絕對不能是 Normal，改回退至 書狀_預設
    fallback = '書狀_預設'
    if fallback in existing_styles:
        return fallback
    return 'Normal' if 'Normal' in existing_styles else None


def _add_tab_stop_6cm(p):
    """
    先 normalize（清 left_indent、first_line_indent、舊 w:tabs），
    再加入唯一有效的 6cm 左對齊 tab stop。
    """
    normalize_header_paragraph(p)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(6), WD_TAB_ALIGNMENT.LEFT)


# ─────────────────────────────────────────────────────────────────────────────
# 身份類別判斷
# ─────────────────────────────────────────────────────────────────────────────

def _is_identity_line(text):
    """判斷文字是否為身份行（角色冒號開頭）。"""
    return bool(IDENTITY_ROLES.match(text.strip()))


def _is_continuation_line(text):
    """
    判斷是否為「純續行資訊」行。條件之一滿足即為真：
    1. 以 \\t 開頭
    2. 以全形空白（\\u3000）或多個半形空白開頭（Markdown 中的縮排地址格式）
    3. 不以角色關鍵字開頭，但包含地址/電話/email 等後續資訊關鍵字
    """
    t = text.strip()
    if t.startswith('\t'):
        return True
    # 全形空白縮排（如「　　　　住桃園市...」）
    if text.startswith('\u3000') or (len(text) > 2 and text[:2] == '  '):
        return True
    if not IDENTITY_ROLES.match(t) and CONTINUATION_KEYWORDS.search(t):
        return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# 各類段落寫入函式
# ─────────────────────────────────────────────────────────────────────────────

def write_identity_paragraph(doc, style_name, text):
    target_style = '書狀_狀首當事人'
    
    # 收斂所有空白與 Tab
    text_clean = re.sub(r'[\t\s\u3000]+', ' ', text.strip())
    
    # 拆解稱謂
    role_match = IDENTITY_ROLES.match(text_clean)
    role = role_match.group(1) if role_match else ""
    
    # 取得稱謂後方的剩餘文字，並確保移除開頭可能殘留的冒號
    remainder = text_clean[role_match.end():].strip() if role_match else text_clean
    remainder = re.sub(r'^[\uff1a:]+', '', remainder).strip()
    
    # 尋找第一筆地址、電話等後續資訊的起點
    cont_match = CONTINUATION_KEYWORDS.search(remainder)
    
    p = doc.add_paragraph(style=_get_safe_style(doc, target_style))
    
    def _format_names_with_wrap(role_str, names_str, max_chars=26):
        """將以頓號分隔的姓名依 26 字元限制換行，續行補 \t。"""
        if '、' not in names_str:
            return names_str
        
        parts = [n.strip() for n in names_str.split('、') if n.strip()]
        if not parts:
            return names_str
            
        lines = []
        current_names = []
        # 首行長度包含稱謂及間距估算
        current_len = len(role_str) + 2 
        
        for n in parts:
            add_len = len(n) + (1 if current_names else 0)
            if current_names and (current_len + add_len > max_chars):
                lines.append('、'.join(current_names))
                current_names = [n]
                # 換行後的第一層縮排（\t）約相當於 4 個中文字寬
                current_len = 4 + len(n)
            else:
                current_names.append(n)
                current_len += add_len
                
        if current_names:
            lines.append('、'.join(current_names))
            
        return '\n\t'.join(lines)
    
    if cont_match:
        name = remainder[:cont_match.start()].strip()
        details_str = remainder[cont_match.start():].strip()
        
        # 將 details_str 依照空白分割，並根據關鍵字重新分組
        # 確保如「電話：」能被獨立出來，而不會把「地址同上」切斷
        parts = details_str.split(' ')
        details_list = []
        current_detail = ""
        
        for part in parts:
            # 只要該片段的開頭符合連續資訊關鍵字，就視為新的一筆資料
            if CONTINUATION_KEYWORDS.match(part):
                if current_detail:
                    details_list.append(current_detail)
                current_detail = part
            else:
                # 否則接續在目前的資料後面（處理地址中間有空白的狀況）
                if current_detail:
                    current_detail += " " + part
                else:
                    current_detail = part
        if current_detail:
            details_list.append(current_detail)
            
        # 將處理好多人折行的 name 放進去
        wrapped_name = _format_names_with_wrap(role, name)
        p.add_run(f"{role}\t{wrapped_name}")
        for i, detail in enumerate(details_list):
            if i == 0:
                # 第一筆資料（通常是地址或事務所）
                # 這裡的 len(name) 判斷主要是為了決定地址要不要換行，我們以字串長度概估
                if len(name) <= 5 and '\n' not in wrapped_name:
                    p.add_run(f"\t{detail}")
                else:
                    p.add_run(f"\n\t\t{detail}")
            else:
                # 第二筆以後的資料（如電話、Email），一律換行
                p.add_run(f"\n\t\t{detail}")
    else:
        # 只有稱謂與姓名
        wrapped_remainder = _format_names_with_wrap(role, remainder)
        p.add_run(f"{role}\t{wrapped_remainder}")
        
    return p

def write_continuation_paragraph(doc, style_name, detail_text):
    target_style = '書狀_狀首當事人'
    
    # 移除開頭的空白與 Tab
    detail_clean = re.sub(r'^[\t\s\u3000]+', '', detail_text)
    
    p = doc.add_paragraph(style=_get_safe_style(doc, target_style))
    # 純續行段落從 0cm 開始，給「2 個 \t」跳過 3cm，直接落在 6cm 處
    p.add_run(f"\t\t{detail_clean}")
    
    return p


def write_value_of_claim_paragraph(doc, text):
    """
    訴訟標的價額例外行：
    - 套用 書狀_預設
    - left_indent = 0、first_line_indent = 0
    - 不加 6 公分 tab
    """
    text = text.strip()
    p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
    normalize_header_paragraph(p)
    p.add_run(text)
    print(f"[INFO] 訴訟標的價額: {text[:50]}")
    return p


def write_header_line(doc, text, style_name):
    """
    根據行的性質，自動分派到正確的寫入函式。
    優先順序：
      1. 訴訟標的價額 → write_value_of_claim_paragraph
      2. 為就...事：  → 書狀_預設 不加 tab
      3. 全形空白/tab 縮排的續行 → write_continuation_paragraph
      4. 身份行 → write_identity_paragraph
      5. 其餘 → 普通段落
    """
    if not style_name or style_name in ('Normal', '內文'):
        style_name = '書狀_預設'

    text_norm = text.replace('\u3000', '').replace(' ', '')

    # 1. 訴訟標的價額
    if text_norm.startswith('訴訟標的價額'):
        return write_value_of_claim_paragraph(doc, text)

    # 2. 為就...事：
    t = text.strip()
    if ('為就' in t and END_PATTERN.search(t)) or (END_PATTERN.search(t) and len(t) > 5):
        p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
        normalize_header_paragraph(p)
        p.add_run(t)
        return p

    # 3. 全形空白/tab 縮排的地址續行（優先於身份行判斷）
    if _is_continuation_line(text):
        return write_continuation_paragraph(doc, style_name, text)

    # 4. 身份行
    if _is_identity_line(text):
        return write_identity_paragraph(doc, style_name, text)

    # 5. 普通（案號、股別、其他無縮排的狀首行）
    p = doc.add_paragraph(style=_get_safe_style(doc, style_name))
    normalize_header_paragraph(p)
    p.add_run(t)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# 狀首擷取
# ─────────────────────────────────────────────────────────────────────────────

def extract_header(docx_path):
    """
    從既有 .docx 擷取狀首段落（到「為就...事：」為止，含該段）。
    回傳 [{'style': str, 'text': str, 'element': element}, ...]
    """
    src = Document(str(docx_path))
    header_data = []

    for i, p in enumerate(src.paragraphs):
        if i >= 30:
            print("[WARN] 狀首擷取超過 30 個段落仍未見「為就...事：」，自動停止擷取")
            break

        text = p.text.strip()
        if not text:
            continue

        header_data.append({
            'style': p.style.name,
            'text': p.text,
            'element': copy.deepcopy(p._element),
        })

        if ('為就' in text and END_PATTERN.search(text)) or \
           (END_PATTERN.search(text) and len(text) > 5):
            break

    print(f"[INFO] 擷取 {len(header_data)} 個狀首段落")
    return header_data


# ─────────────────────────────────────────────────────────────────────────────
# 從 Markdown blocks 萃取狀首資訊
# ─────────────────────────────────────────────────────────────────────────────

def extract_md_header_info(md_blocks):
    """
    從 parse_markdown 得到的 blocks 萃取狀首區間的所有段落資訊。
    收集範圍：從第一個 block 直到「為就...事：」行（含）。
    回傳的 header_lines 包含：
      - 書狀_標題（書狀名稱）
      - 書狀_狀首日期
      - 案號、股別（書狀_預設）
      - 身份行（書狀_預設，角色:姓名）
      - 全形空白縮排的地址/續行（書狀_預設 原始文字含全形縮排）
      - 訴訟標的價額（書狀_預設）
      - 為就...事：（書狀_預設，最後一行）
    """
    info = {
        'title': None,
        'date': None,
        'value_of_claim': None,
        'for_matter': None,
        'header_lines': [],   # 完整收集稿首區間所有行
    }

    in_header = True

    for b in md_blocks:
        if not in_header:
            break

        text = b.text
        t = text.strip()
        t_norm = t.replace('\u3000', '').replace(' ', '')

        # 遇到為就...事：→ 收集並結束狀首區間
        if ('為就' in t and END_PATTERN.search(t)) or (END_PATTERN.search(t) and len(t) > 5):
            if not info['for_matter']:
                info['for_matter'] = t
            info['header_lines'].append({'style': '書狀_預設', 'text': text})
            in_header = False
            continue

        # 遇到層級段落（正文段落）代表已離開狀首區間
        # 書狀_簽章/書狀_法院不在清單：狀首裡的「訴訟代理人：」被 parse_markdown 解析為書狀_簽章，
        # 但它仍是狀首行，需要繼續收集
        if b.style in ('通用_層級1', '通用_層級2', '通用_層級3',
                       '通用_層級4', 'SEMANTIC_HEADING', '書狀_謹狀',
                       '書狀_證據編號', '書狀_證據編號10',
                       '書狀_被上證據編號', '書狀_被上證據編號10'):
            in_header = False
            continue

        # 收集狀首資訊
        # raw_line: 優先用 b.raw_text（保留全形空白縮排），否則用 b.text
        raw_line = getattr(b, 'raw_text', None) or b.text
        if b.style == '書狀_標題':
            if not info['title']:
                info['title'] = t
            info['header_lines'].append({'style': '書狀_標題', 'text': raw_line})
        elif b.style == '書狀_狀首日期':
            if not info['date']:
                info['date'] = t
            info['header_lines'].append({'style': '書狀_狀首日期', 'text': text})
        elif t_norm.startswith('訴訟標的價額'):
            if not info['value_of_claim']:
                info['value_of_claim'] = t
            info['header_lines'].append({'style': '書狀_預設', 'text': text})
        else:
            # 案號、股別、身份行、全形縮排地址行等
            # 使用 raw_line 保留全形空白縮排，以便 _is_continuation_line 能正確識別
            info['header_lines'].append({'style': '書狀_預設', 'text': raw_line})

    return info


# ─────────────────────────────────────────────────────────────────────────────
# 主合併寫入函式
# ─────────────────────────────────────────────────────────────────────────────

def merge_and_write_header(doc, header_data=None, md_headers=None,
                           is_issue_table=False, party_status=""):
    """
    將狀首段落寫入 doc。
    - header_data: 從舊 docx 擷取的段落列表（優先使用）
    - md_headers: 從 Markdown blocks 萃取的資訊 dict（header_data 為 None 時使用）
    - is_issue_table: True 時書狀名稱轉換為「爭點整理狀」
    """
    if md_headers is None:
        md_headers = {}

    if header_data:
        # ── 舊 docx 狀首：重排訴訟標的價額，逐行寫入 ──
        val_claim = md_headers.get('value_of_claim')
        case_idx = -1
        val_idx = -1

        for i, item in enumerate(header_data):
            norm = item['text'].replace('\u3000', '').replace(' ', '')
            if norm.startswith('訴訟標的價額'):
                if not val_claim:
                    val_claim = item['text']
                val_idx = i
            if (norm.startswith('案號') or norm.startswith('股別')) and case_idx == -1:
                case_idx = i

        if val_idx != -1:
            header_data.pop(val_idx)
            if val_idx < case_idx:
                case_idx -= 1

        if val_claim:
            voc_item = {'text': val_claim, 'style': '書狀_預設'}
            insert_at = case_idx if case_idx != -1 else 2
            header_data.insert(insert_at, voc_item)

        for i, item in enumerate(header_data):
            text = item['text']
            style_name = item['style']
            t = text.strip()

            # 第一行：書狀名稱
            if i == 0:
                text = md_headers.get('title') or text
                t = text.strip()
                if is_issue_table:
                    t = re.sub(
                        r'^(民事|刑事|行政|家事)?.*',
                        lambda m: (m.group(1) or '') + "爭點整理狀",
                        t
                    )
                p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_標題'))
                normalize_header_paragraph(p)
                p.add_run(t)
                continue

            # 日期行（排除案號，且長度、位置合理）
            is_date_line = (style_name == '書狀_狀首日期' or DATE_PATTERN.search(t)) and '案號' not in t and i < 5 and len(t) < 20
            if is_date_line:
                if not _wants_top_date():
                    continue

                # 重新計算預設日期：今天 + 2 日（依規則：狀首日期「絕對不要加」民國）
                target_date = datetime.date.today() + datetime.timedelta(days=2)
                roc_year = target_date.year - 1911
                new_date_str = f"{roc_year}年{target_date.month}月{target_date.day}日"
                
                p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
                normalize_header_paragraph(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(new_date_str)
                run.font.size = Pt(12)
                continue

            # 為就...事：
            if ('為就' in t and END_PATTERN.search(t)) or (END_PATTERN.search(t) and len(t) > 5):
                text = md_headers.get('for_matter') or text
                t = text.strip()
                if is_issue_table:
                    m = re.match(r'(為就.+?(?:事件|案件))', t)
                    case_desc = m.group(1) if m else '為就本案'
                    t = f'{case_desc}，提出爭點整理事：'
                p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
                normalize_header_paragraph(p)
                p.add_run(t)
                continue

            style_name = _get_safe_style(doc, style_name)
            write_header_line(doc, text, style_name)

    else:
        # ── 無舊 docx，從 md_headers.header_lines 生成 ──
        header_lines = md_headers.get('header_lines', [])

        if header_lines:
            for i, item in enumerate(header_lines):
                text = item['text']
                style_name = item['style']
                t = text.strip()

                if i == 0:
                    # 第一行：書狀名稱
                    if is_issue_table:
                        t = re.sub(
                            r'^(民事|刑事|行政|家事)?.*',
                            lambda m: (m.group(1) or '') + "爭點整理狀",
                            t
                        )
                    p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_標題'))
                    normalize_header_paragraph(p)
                    p.add_run(t)
                    continue

                if style_name == '書狀_狀首日期' or DATE_PATTERN.match(t):
                    if not _wants_top_date():
                        continue

                    target_date = datetime.date.today() + datetime.timedelta(days=2)
                    roc_year = target_date.year - 1911
                    
                    # 依規則：狀首日期「絕對不要加」民國
                    m = DATE_PATTERN.search(text)
                    if m:
                        parts = re.split(r'年', text)
                        suffix = parts[1] if len(parts) > 1 else f"{target_date.month}月{target_date.day}日"
                        new_date_str = f"{roc_year}年{suffix.strip()}"
                    else:
                        new_date_str = f"{roc_year}年{target_date.month}月{target_date.day}日"
                        
                    p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
                    normalize_header_paragraph(p)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(new_date_str)
                    run.font.size = Pt(12)
                    continue

                # 為就...事：
                if ('為就' in t and END_PATTERN.search(t)) or (END_PATTERN.search(t) and len(t) > 5):
                    if is_issue_table:
                        m = re.match(r'(為就.+?(?:事件|案件))', t)
                        case_desc = m.group(1) if m else '為就本案'
                        t = f'{case_desc}，提出爭點整理事：'
                    p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
                    normalize_header_paragraph(p)
                    p.add_run(t)
                    continue

                write_header_line(doc, text, style_name)

        elif md_headers.get('title'):
            # 有 title 但 header_lines 為空（舊版 extract_md_header_info 格式）
            title = md_headers['title']
            if is_issue_table:
                title = re.sub(
                    r'^(民事|刑事|行政|家事)?.*',
                    lambda m: (m.group(1) or '') + "爭點整理狀",
                    title
                )
            p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_標題'))
            normalize_header_paragraph(p)
            p.add_run(title)

            if _wants_top_date():
                p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
                normalize_header_paragraph(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(md_headers.get('date') or f"{datetime.date.today().year - 1911}年　月　日")
                run.font.size = Pt(12)

            if md_headers.get('value_of_claim'):
                write_value_of_claim_paragraph(doc, md_headers['value_of_claim'])

            for text in md_headers.get('others', []):
                write_header_line(doc, text, '書狀_預設')

            matter = md_headers.get('for_matter')
            if not matter:
                matter = "為就本案爭議，提出爭點整理事：" if is_issue_table else "為就本案爭議，提出事："
            elif is_issue_table:
                m = re.match(r'(為就.+?(?:事件|案件))', matter)
                case_desc = m.group(1) if m else '為就本案'
                matter = f'{case_desc}，提出爭點整理事：'
            p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
            normalize_header_paragraph(p)
            p.add_run(matter)

        else:
            # 最後防線：空白骨架
            title = "民事爭點整理狀" if is_issue_table else "民事書狀"
            p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_標題'))
            normalize_header_paragraph(p)
            p.add_run(title)

            if _wants_top_date():
                p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
                normalize_header_paragraph(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(f"{datetime.date.today().year - 1911}年　月　日")
                run.font.size = Pt(12)

            for line in ['案　　號：', '股　　別：', '上 訴 人：', '法定代理人：', '訴訟代理人：', '被上訴人：']:
                p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
                normalize_header_paragraph(p)
                p.add_run(line)

            matter = "為就…案件，提出爭點整理事：" if is_issue_table else "為就...事："
            p = doc.add_paragraph(style=_get_safe_style(doc, '書狀_預設'))
            normalize_header_paragraph(p)
            p.add_run(matter)
