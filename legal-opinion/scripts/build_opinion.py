#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_opinion.py — 法律意見書 Markdown 轉 DOCX 排版引擎

使用 python-docx 載入範本，執行錨點反查、清除內容、解析 Markdown，
並依據範本實際樣式套用「法律書函_標題」、「法律書函_預設」、「法律書函_主旨」、
「通用_層級1~4」、「法律書函_附件」、「法律書函_附件10」、「法律書函_簽章」。
支援日期轉換為公元紀年並靠右 12pt 對齊。
"""

import sys
import os
import re
import argparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Twips, Pt
from lxml import etree

# 確保 stdout 使用 UTF-8
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

# ==============================================================================
# 常數與規則定義
# ==============================================================================

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# 錨點預設樣式名稱（與範本中的樣式一致）
ANCHOR_STYLE = '通用_層級1'

BODY_INDENT_MAP = {
    0: 567,   # 通用_層級1 下方的論述段落 (10 mm)
    1: 850,   # 通用_層級2 下方的論述段落 (15 mm)
    2: 850,   # 通用_層級3 下方的論述段落
    3: 1134,  # 通用_層級4 下方的論述段落 (20 mm)
}

# Markdown 大綱前綴匹配
LEVEL_PATTERNS = [
    (re.compile(r'^([一二三四五六七八九十百千]+)、\s*(.*)'), '通用_層級1', 0),
    (re.compile(r'^[\(（]([一二三四五六七八九十百千]+)[\)）]\s*(.*)'), '通用_層級2', 1),
    (re.compile(r'^(\d+)\.\s+(.*)'), '通用_層級3', 2),
    (re.compile(r'^\((\d+)\)\s*(.*)'), '通用_層級4', 3),
]

# 其它特殊段落匹配
DATE_PATTERN = re.compile(r'^日期\s*[：:]\s*(.*)')
RECEIVER_PATTERN = re.compile(r'^受文者\s*[：:]\s*(.*)')
SUBJECT_PATTERN = re.compile(r'^主旨\s*[：:]\s*(.*)')
EXPLANATION_PATTERN = re.compile(r'^說明\s*[：:]\s*(.*)')
ATTACHMENT_HEADER_PATTERN = re.compile(r'^附件\s*[：:]\s*(.*)')
ATTACHMENT_ITEM_PATTERN = re.compile(r'^附件\s*(\d+)\s*[：:]\s*(.*)')
SIGNATURE_PATTERN = re.compile(r'^撰寫人\s*[：:]\s*(.*)')

BOLD_PATTERN = re.compile(r'\*\*(.+?)\*\*')

# ==============================================================================
# 區塊資料模型
# ==============================================================================

class Block:
    def __init__(self, style, text, ilvl=None, needs_num=False, is_date=False, is_attachment=False, is_subject=False):
        self.style = style
        self.text = text
        self.ilvl = ilvl               # 0~3
        self.needs_num = needs_num
        self.is_date = is_date
        self.is_attachment = is_attachment
        self.is_subject = is_subject

# ==============================================================================
# 日期公元換算
# ==============================================================================

def convert_to_ad_date(date_str):
    """將日期字串（支援民國或其它格式）換算並格式化為公元紀年。"""
    date_str = date_str.strip()
    # 移除字串內多餘空格
    cleaned = re.sub(r'\s+', '', date_str)
    
    # 匹配 中華民國YY年MM月DD日 或 民國YY年MM月DD日 或 YY年MM月DD日
    m_roc = re.match(r'^(?:中華民國|民國)?(\d+)年(\d+)月(\d+)日$', cleaned)
    if m_roc:
        year = int(m_roc.group(1))
        month = int(m_roc.group(2))
        day = int(m_roc.group(3))
        # 若年份小於 1000，視為民國紀年，換算為公元
        if year < 1000:
            year += 1911
        return f"日期：{year}年{month}月{day}日"
        
    # 匹配 YYYY-MM-DD 或 YYYY/MM/DD 等
    m_dash = re.match(r'^(\d{4})[-/](\d{1,2})[-/](\d{1,2})$', cleaned)
    if m_dash:
        year = int(m_dash.group(1))
        month = int(m_dash.group(2))
        day = int(m_dash.group(3))
        return f"日期：{year}年{month}月{day}日"
        
    # 若無匹配則僅加上前綴回傳
    if not date_str.startswith("日期：") and not date_str.startswith("日期:"):
        return f"日期：{date_str}"
    return date_str

# ==============================================================================
# 核心排版邏輯
# ==============================================================================

def extract_subject_ppr_overrides(doc):
    """在清空 Body 之前，從範例檔中提取主旨段落的段落層級覆寫屬性。
    
    範例檔中「法律書函_主旨」段落的縮排 (ind) 和編號 (numPr)
    定義在段落層級而非樣式定義中。若不提取並複製，新段落將遺失這些屬性。
    回傳：pPr 覆寫子節點清單（不含 pStyle），若無則回傳空清單。
    """
    SUBJECT_STYLE = '法律書函_主旨'
    overrides = []
    for p in doc.paragraphs:
        if p.style.name == SUBJECT_STYLE:
            pPr = p._element.find(qn('w:pPr'))
            if pPr is not None:
                from copy import deepcopy
                for child in pPr:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    # 排除 pStyle（已由樣式套用）和 overflowPunct（由 disable_hanging_punctuation 處理）
                    if tag not in ('pStyle', 'overflowPunct'):
                        overrides.append(deepcopy(child))
            break  # 只取第一個主旨段落
    return overrides

def find_and_remove_anchor(doc):
    """反查錨點樣式對應的編號 ID。"""
    anchor_para = None
    for p in doc.paragraphs:
        if p.style.name == ANCHOR_STYLE:
            anchor_para = p
            break

    if anchor_para is None:
        raise RuntimeError(f"模板中找不到 '{ANCHOR_STYLE}' 樣式的錨點段落")

    pPr = anchor_para._element.get_or_add_pPr()
    num_id = None
    if pPr.numPr is not None and pPr.numPr.numId is not None:
        num_id = pPr.numPr.numId.val
    else:
        num_id = _trace_num_id_from_style(doc, anchor_para.style)

    abstract_num_id = _get_abstract_num_id(doc, num_id)
    anchor_para._element.getparent().remove(anchor_para._element)
    return num_id, abstract_num_id

def _trace_num_id_from_style(doc, style):
    curr = style
    while curr:
        pPr = curr._element.find('.//w:pPr', NSMAP)
        if pPr is not None:
            numPr = pPr.find('w:numPr', NSMAP)
            if numPr is not None:
                nid = numPr.find('w:numId', NSMAP)
                if nid is not None: 
                    return int(nid.get(qn('w:val')))
        if curr.base_style: 
            curr = curr.base_style
        else: 
            break
    raise RuntimeError("無法從樣式鏈取得編號 ID")

def _get_abstract_num_id(doc, num_id):
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    for num in numbering.findall('.//w:num', NSMAP):
        if int(num.get(qn('w:numId'))) == num_id:
            return int(num.find('w:abstractNumId', NSMAP).get(qn('w:val')))
    raise RuntimeError(f"找不到對應的 abstractNumId (numId={num_id})")

def create_override_num(doc, abstract_num_id):
    """建立新的編號實例以重新起算編號。"""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    max_id = max([int(n.get(qn('w:numId'))) for n in numbering.findall('.//w:num', NSMAP)] + [0])
    new_id = max_id + 1
    
    num_el = etree.SubElement(numbering, qn('w:num'))
    num_el.set(qn('w:numId'), str(new_id))
    etree.SubElement(num_el, qn('w:abstractNumId')).set(qn('w:val'), str(abstract_num_id))
    
    for lvl in range(4):
        ov = etree.SubElement(num_el, qn('w:lvlOverride'))
        ov.set(qn('w:ilvl'), str(lvl))
        etree.SubElement(ov, qn('w:startOverride')).set(qn('w:val'), '1')
    return new_id

def setup_page_rules(doc):
    """設定行編號。"""
    for sec in doc.sections:
        sectPr = sec._sectPr
        ln = sectPr.find(qn('w:lnNumType'))
        if ln is None:
            ln = OxmlElement('w:lnNumType')
            sectPr.append(ln)
        ln.set(qn('w:countBy'), '1')
        ln.set(qn('w:restart'), 'newPage')

def disable_hanging_punctuation(p):
    """關閉段落中文標點懸尾。"""
    pPr = p._element.get_or_add_pPr()
    for el in pPr.findall(qn('w:overflowPunct')): 
        pPr.remove(el)
    etree.SubElement(pPr, qn('w:overflowPunct')).set(qn('w:val'), '0')

def apply_body_indent(p, ilvl):
    """套用論述段落左側縮排。"""
    indent = BODY_INDENT_MAP.get(ilvl, 0)
    if indent > 0:
        p.paragraph_format.left_indent = Twips(indent)

# ==============================================================================
# 解析與寫入
# ==============================================================================

def parse_markdown(path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    blocks = []
    in_front_matter = False
    
    for line in lines:
        line = line.strip()
        
        # 處理 YAML Front Matter 區塊
        if line == '---':
            in_front_matter = not in_front_matter
            continue
        if in_front_matter:
            continue
            
        if not line: 
            continue
        
        # 1. 標題 Heading 1 & 2
        if line.startswith('# '):
            text = line[2:].strip()
            # 若為「法律意見書」或「# 法律意見書」，套用標題樣式
            if "法律意見書" in text:
                blocks.append(Block('法律書函_標題', text))
            else:
                blocks.append(Block('法律書函_標題', text))
            continue
        if line.startswith('## '):
            blocks.append(Block('法律書函_標題', line[3:].strip()))
            continue
            
        # 2. 特殊關鍵行識別
        m_date = DATE_PATTERN.match(line)
        if m_date:
            raw_val = m_date.group(1)
            converted = convert_to_ad_date(raw_val)
            blocks.append(Block('法律書函_預設', converted, is_date=True))
            continue
            
        if RECEIVER_PATTERN.match(line):
            blocks.append(Block('法律書函_預設', line))
            continue
            
        if SUBJECT_PATTERN.match(line):
            blocks.append(Block('法律書函_主旨', line, is_subject=True))
            continue
            
        if EXPLANATION_PATTERN.match(line):
            blocks.append(Block('法律書函_預設', line))
            continue
            
        if ATTACHMENT_HEADER_PATTERN.match(line):
            blocks.append(Block('法律書函_預設', line))
            continue
            
        m_attach = ATTACHMENT_ITEM_PATTERN.match(line)
        if m_attach:
            num = int(m_attach.group(1))
            style = '法律書函_附件' if num < 10 else '法律書函_附件10'
            blocks.append(Block(style, line, is_attachment=True))
            continue
            
        if SIGNATURE_PATTERN.match(line):
            blocks.append(Block('法律書函_簽章', line))
            continue
            
        # 3. 大綱前綴匹配
        matched = False
        for pattern, style, ilvl in LEVEL_PATTERNS:
            m = pattern.match(line)
            if m:
                # 這裡要將前綴拿掉，利用 Word 自動編號
                blocks.append(Block(style, m.group(2).strip(), ilvl=ilvl, needs_num=True))
                matched = True
                break
        if matched: 
            continue
        
        # 4. 預設 法律書函_預設
        blocks.append(Block('法律書函_預設', line))
        
    return blocks

def write_block(doc, block, num_id, subject_overrides=None):
    # 檢查 style 是否在模板中，不存在則回退至 Normal
    style_name = block.style
    available_styles = [s.name for s in doc.styles]
    if style_name not in available_styles:
        style_name = 'Normal'
        
    p = doc.add_paragraph(style=style_name)
    text = block.text
    
    # 處理字元粗體樣式
    parts = []
    last = 0
    for m in BOLD_PATTERN.finditer(text):
        if m.start() > last: 
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text): 
        parts.append((text[last:], False))
    
    for t_part, is_bold in (parts if parts else [(text, False)]):
        run = p.add_run(t_part)
        if is_bold: 
            run.bold = True
            
    # 特殊處理日期：靠右
    if block.is_date:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 主旨段落：複製範例檔中的段落層級覆寫（numPr、ind 等）
    if block.is_subject and subject_overrides:
        pPr = p._element.get_or_add_pPr()
        from copy import deepcopy
        for override in subject_overrides:
            pPr.append(deepcopy(override))
            
    # 設定編號屬性 (numPr)
    if block.needs_num and block.ilvl is not None:
        pPr = p._element.get_or_add_pPr()
        for old in pPr.findall(qn('w:numPr')): 
            pPr.remove(old)
        numPr = etree.SubElement(pPr, qn('w:numPr'))
        etree.SubElement(numPr, qn('w:ilvl')).set(qn('w:val'), str(block.ilvl))
        etree.SubElement(numPr, qn('w:numId')).set(qn('w:val'), str(num_id))
    
    disable_hanging_punctuation(p)
    return p

# ==============================================================================
# 主程式
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(description='法律意見書排版引擎')
    parser.add_argument('draft', help='Markdown 草稿路徑')
    parser.add_argument('--template', help='Word 模板路徑')
    parser.add_argument('--output', help='輸出 DOCX 路徑')
    args = parser.parse_args()

    # 預設模板路徑
    tpl_path = args.template or os.path.join(os.path.dirname(__file__), '..', 'assets', 'opinion-tmpl.docx')
    if not os.path.exists(tpl_path):
        print(f"[ERROR] 找不到模板檔案: {tpl_path}", file=sys.stderr)
        sys.exit(1)
        
    doc = Document(tpl_path)
    
    # 1. 初始化：提取主旨覆寫屬性（必須在清空 Body 前執行）
    subject_overrides = extract_subject_ppr_overrides(doc)
    anchor_num_id, abstract_num_id = find_and_remove_anchor(doc)
    
    body = doc.element.body
    for c in list(body):
        if c.tag != qn('w:sectPr'): 
            body.remove(c)
            
    setup_page_rules(doc)
    
    # 2. 解析 Markdown
    blocks = parse_markdown(args.draft)
    
    # 3. 寫入
    curr_num_id = anchor_num_id
    last_ilvl = None
    
    # 尋找特定 Block 的索引以插入空行
    first_attachment_idx = None
    signature_idx = None
    for idx, b in enumerate(blocks):
        if b.style == '法律書函_簽章':
            signature_idx = idx
        if (b.is_attachment or b.text.strip().startswith("附件：") or b.text.strip().startswith("附件:")) and first_attachment_idx is None:
            first_attachment_idx = idx
    
    for idx, b in enumerate(blocks):
        # (1) 有附件時，附件和本文（之前段落）之間要空一行
        if first_attachment_idx is not None and idx == first_attachment_idx:
            p_empty = doc.add_paragraph(style='Normal')
            disable_hanging_punctuation(p_empty)
            
        # (2) 撰寫人（簽章）之前要空一行（不論有無附件）
        if idx == signature_idx:
            p_empty = doc.add_paragraph(style='Normal')
            disable_hanging_punctuation(p_empty)

        # 當出現新的第一層大綱，重設編號 instance
        if b.style == '通用_層級1' and b.ilvl == 0:
            curr_num_id = create_override_num(doc, abstract_num_id)
            
        p = write_block(doc, b, curr_num_id, subject_overrides)
        
        # 論述縮排
        if b.style == '法律書函_預設' and last_ilvl is not None:
            apply_body_indent(p, last_ilvl)
            
        # 追蹤最後大綱層級
        if b.needs_num: 
            last_ilvl = b.ilvl
        elif b.style != '法律書函_預設': 
            last_ilvl = None
        
    out = args.output or (os.path.splitext(args.draft)[0] + "_output.docx")
    doc.save(out)
    print(f"成功產出 Word 檔: {out}")

if __name__ == "__main__":
    main()
