# -*- coding: utf-8 -*-
# postprocess_const_pleading.py

import sys
import os
import re
import argparse
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm

# 強制標準輸出採用 UTF-8，避免 Windows 環境印出中文字元出錯
sys.stdout.reconfigure(encoding='utf-8')

# 設定 NSMAP
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
BLOCK_HEADING_STYLE = '書狀_區塊標題'


def ensure_block_heading_style(doc):
    """Create the constitution-only block heading style if the template does not have it."""
    try:
        return doc.styles[BLOCK_HEADING_STYLE]
    except KeyError:
        pass

    style = doc.styles.add_style(BLOCK_HEADING_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    try:
        style.base_style = doc.styles['書狀_預設']
    except KeyError:
        pass
    style.paragraph_format.left_indent = Cm(0)
    style.paragraph_format.first_line_indent = Cm(0)
    return style


def strip_first_text_run_prefix(paragraph):
    for run in paragraph.runs:
        if run.text:
            run.text = re.sub(r'^[　\s*_]+', '', run.text)
            return


def prefix_first_text_run(paragraph, prefix):
    for run in paragraph.runs:
        if run.text is not None:
            run.text = prefix + run.text
            return

def main():
    parser = argparse.ArgumentParser(description="憲法書狀 DOCX 格式後處理器")
    parser.add_argument("input_docx", help="輸入的 DOCX 檔案路徑")
    parser.add_argument("--output", help="輸出的 DOCX 檔案路徑 (若未指定則直接覆寫輸入檔)")
    args = parser.parse_args()

    input_path = os.path.abspath(args.input_docx)
    output_path = os.path.abspath(args.output) if args.output else input_path

    if not os.path.exists(input_path):
        print(f"[ERROR] 找不到輸入檔案: {input_path}")
        sys.exit(1)

    print(f"[INFO] 開始處理憲法書狀後處理: {input_path}")
    doc = Document(input_path)
    ensure_block_heading_style(doc)

    # 動態引入 draft-pleading 中的 word_xml_utils.py
    script_dir = os.path.dirname(os.path.abspath(__file__))
    skills_dir = os.path.abspath(os.path.join(script_dir, "..", ".."))
    draft_pleading_scripts = os.path.join(skills_dir, "draft-pleading", "scripts")
    if draft_pleading_scripts not in sys.path:
        sys.path.append(draft_pleading_scripts)

    try:
        from word_xml_utils import (
            _get_abstract_num_id, _trace_num_id_from_style,
            create_override_num, create_l2_reset_num, set_num_pr
        )
    except ImportError as e:
        print(f"[ERROR] 無法從 {draft_pleading_scripts} 載入 word_xml_utils: {e}")
        sys.exit(1)

    # 1. 找出 abstractNumId
    abstract_num_id = None
    # 嘗試從已套用編號的段落中尋找
    for p in doc.paragraphs:
        pPr = p._element.pPr
        if pPr is not None:
            numPr = pPr.find('.//w:numPr', NSMAP)
            if numPr is not None:
                nid_el = numPr.find('w:numId', NSMAP)
                if nid_el is not None:
                    try:
                        num_id = int(nid_el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
                        abstract_num_id = _get_abstract_num_id(doc, num_id)
                        break
                    except Exception:
                        pass

    # 備用方案：從樣式中反查
    if abstract_num_id is None:
        try:
            if '通用_層級1' in doc.styles:
                style = doc.styles['通用_層級1']
                num_id = _trace_num_id_from_style(doc, style)
                abstract_num_id = _get_abstract_num_id(doc, num_id)
        except Exception as e:
            print(f"[WARNING] 無法從樣式反查 abstractNumId: {e}")

    if abstract_num_id is None:
        print("[ERROR] 無法取得文件的 abstractNumId，無法重設編號。")
        sys.exit(1)

    print(f"[INFO] 取得 abstractNumId: {abstract_num_id}")

    # 2. 定義標題比對規則
    cjk_num_pat = re.compile(r'^[壹貳參肆伍陸柒捌玖拾]+[、]')
    semantic_headings = {
        '聲請審查客體', '應受判決事項之聲明', '主要爭點', '聲請理由', '理由', '事實與理由'
    }

    # 3. 預先掃描，找出有子段落的一級段落
    has_children = set()
    current_l1_idx = None
    for idx, p in enumerate(doc.paragraphs):
        sname = p.style.name
        if sname in ('通用_層級1', '通用_5f_層級1'):
            current_l1_idx = idx
        elif sname in ('通用_層級2', '通用_5f_層級2', '通用_層級3', '通用_5f_層級3', '通用_層級4', '通用_5f_層級4'):
            if current_l1_idx is not None:
                has_children.add(current_l1_idx)
        elif p.text.strip():
            # 遇到其他大標題，重設一級段落定位，避免跨區判定
            ct = p.text.strip().replace('\u3000', '').replace(' ', '').replace('*', '').replace('_', '')
            if ct in semantic_headings or cjk_num_pat.match(ct):
                current_l1_idx = None

    # 4. 開始掃描並進行格式後處理
    reset_pending = False
    bold_section = False
    current_l1_num_id = None
    _last_l1_num_id = None
    l2_cache = {}

    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        clean_text = text.replace('\u3000', '').replace(' ', '').replace('*', '').replace('_', '')
        sname = p.style.name

        # A. 處理語意標題
        if clean_text in semantic_headings:
            # 使用專屬區塊標題樣式承載兩個全形空白，不刪除底層 w:ind。
            # 直接移除縮排 XML 會讓 LibreOffice 轉 ODT 時失去清單層級與重設線索。
            try:
                p.style = BLOCK_HEADING_STYLE
            except Exception:
                pass

            strip_first_text_run_prefix(p)
            prefix_first_text_run(p, "　　")
            
            reset_pending = True

        # B. 處理大寫國字數字標題
        elif cjk_num_pat.match(clean_text):
            # 壹、貳等大類別同樣使用區塊標題樣式靠左呈現，不直接刪 XML 縮排節點。
            try:
                p.style = BLOCK_HEADING_STYLE
            except Exception:
                pass
            strip_first_text_run_prefix(p)
            
            # 判斷是否進入違憲理由區間
            if '違憲理由' in clean_text:
                bold_section = True
            else:
                bold_section = False
            
            reset_pending = True

        # C. 處理通用層級段落
        elif sname in ('通用_層級1', '通用_5f_層級1'):
            if reset_pending:
                # 建立新的一級 numId，以重設編號為 1
                current_l1_num_id = create_override_num(doc, abstract_num_id)
                _last_l1_num_id = current_l1_num_id
                reset_pending = False

            if current_l1_num_id is not None:
                set_num_pr(p, current_l1_num_id, ilvl=0, outline_level=0)

            # 處理粗體：若在違憲理由區間內，且該一級段落有子段落
            if bold_section and idx in has_children:
                for run in p.runs:
                    run.font.bold = True

        elif sname in ('通用_層級2', '通用_5f_層級2'):
            if _last_l1_num_id is not None:
                # 重設二級編號
                if _last_l1_num_id not in l2_cache:
                    l2_cache[_last_l1_num_id] = create_l2_reset_num(doc, abstract_num_id, _last_l1_num_id)
                l2_num_id = l2_cache[_last_l1_num_id]
                set_num_pr(p, l2_num_id, ilvl=1, outline_level=1)

        elif sname in ('通用_層級3', '通用_5f_層級3'):
            if _last_l1_num_id is not None:
                if _last_l1_num_id not in l2_cache:
                    l2_cache[_last_l1_num_id] = create_l2_reset_num(doc, abstract_num_id, _last_l1_num_id)
                l2_num_id = l2_cache[_last_l1_num_id]
                set_num_pr(p, l2_num_id, ilvl=2, outline_level=2)

        elif sname in ('通用_層級4', '通用_5f_層級4'):
            if _last_l1_num_id is not None:
                if _last_l1_num_id not in l2_cache:
                    l2_cache[_last_l1_num_id] = create_l2_reset_num(doc, abstract_num_id, _last_l1_num_id)
                l2_num_id = l2_cache[_last_l1_num_id]
                set_num_pr(p, l2_num_id, ilvl=3, outline_level=3)

    # 5. 合併訴訟代理人與事務所地址（防止雲端同步延遲或 MD 解析分段）
    idx = 0
    while idx < len(doc.paragraphs) - 1:
        p_curr = doc.paragraphs[idx]
        p_next = doc.paragraphs[idx + 1]
        
        text_curr = p_curr.text.strip()
        text_next = p_next.text.strip()
        
        if '訴訟代理人' in text_curr and '張清浩' in text_curr:
            clean_next = text_next.replace('\u3000', '').replace(' ', '').replace('*', '').replace('_', '')
            if clean_next.startswith('事務所：') or '事務所：' in clean_next:
                # 執行合併，使用定位字元 \t 以對齊到 6 公分的定位點
                p_curr.add_run('\t')
                
                # 取得乾淨的事務所地址文字（剝除開頭的空白與定位字元）
                next_full_text = p_next.text.strip()
                clean_address = re.sub(r'^[　\s*\t_]+', '', next_full_text)
                
                # 新增乾淨的地址 run
                new_run = p_curr.add_run(clean_address)
                
                # 繼承字型格式
                if p_next.runs:
                    first_run = p_next.runs[0]
                    new_run.bold = first_run.bold
                    new_run.italic = first_run.italic
                    new_run.font.name = first_run.font.name
                    new_run.font.size = first_run.font.size
                
                # 從 doc 中移除 p_next 段落
                p_next._element.getparent().remove(p_next._element)
                print(f"[INFO] 成功合併訴訟代理人與事務所地址 (Tab對齊): {p_curr.text!r}")
                # 移除了段落，不增加 idx，繼續檢查同一位置（可能是電話）
                continue
        idx += 1

    # 6. 優化已在同一行的訴訟代理人與事務所地址（將姓名與地址之間的空白改為 \t 定位點）
    for p in doc.paragraphs:
        if p.style.name == '書狀_狀首當事人' and '訴訟代理人' in p.text and '事務所：' in p.text:
            # 尋找第一個含有 '事務所：' 的 run
            for r_idx, run in enumerate(p.runs):
                if '事務所：' in run.text:
                    # 清除當前 run 開頭的空白與定位點
                    run.text = re.sub(r'^[　\s*\t_]+', '', run.text)
                    
                    # 向左清理前面所有 runs 的尾端空白
                    for prev_idx in range(r_idx - 1, -1, -1):
                        prev_run = p.runs[prev_idx]
                        if prev_run.text:
                            prev_run.text = re.sub(r'[　\s*\t_]+$', '', prev_run.text)
                            break
                    
                    # 在當前 run 的開頭加上 \t 以進行 6cm 定位
                    run.text = '\t' + run.text
                    print(f"[INFO] 成功優化同行訴訟代理人地址定位 (Tab對齊): {p.text!r}")
                    break

    doc.save(output_path)
    print(f"[INFO] 憲法書狀後處理完成。儲存檔案至: {output_path}")

if __name__ == "__main__":
    main()
