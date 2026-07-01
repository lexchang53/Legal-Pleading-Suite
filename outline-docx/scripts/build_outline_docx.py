#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_outline_docx.py — 通用多層次大綱排版引擎 (v3)

使用 python-docx 載入模板，執行錨點反查、Markdown 解析與動態編號重設，
產出帶有四層中文/阿拉伯數字編號與懸掛縮排的 Word 檔。

v3 更新：
  - 簡化 Markdown 解析，僅支援 Heading 1/2 與四層大綱前綴。
  - Fallback 樣式回歸為 'Normal'。
  - 保留核心的多層次編號 Override 邏輯與論述縮排機制。
"""

import sys
import os
import re
import argparse
import copy
from docx import Document
from docx.oxml.ns import qn, nsmap as oxml_nsmap
from docx.oxml import OxmlElement
from docx.shared import Twips
from lxml import etree

# 確保 stdout 使用 UTF-8
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

# ==============================================================================
# 常數與規則定義
# ==============================================================================

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# 錨點預設樣式名稱
ANCHOR_STYLE = '通用_層級1'

# 層級樣式與縮排設定
LEVEL_STYLES = {
    '通用_層級1': 0,
    '通用_層級2': 1,
    '通用_層級3': 2,
    '通用_層級4': 3,
}

BODY_INDENT_MAP = {
    0: 567,   # 一、 下方的論述段落
    1: 850,   # (一) 下方的論述段落
    2: 850,   # 1.   下方的論述段落
    3: 1134,  # (1)  下方的論述段落
}

# Markdown 前綴匹配
LEVEL_PATTERNS = [
    (re.compile(r'^([一二三四五六七八九十百千]+)、\s*(.*)'), '通用_層級1', 0),
    (re.compile(r'^[\(（]([一二三四五六七八九十百千]+)[\)）]\s*(.*)'), '通用_層級2', 1),
    (re.compile(r'^(\d+)\.\s+(.*)'), '通用_層級3', 2),
    (re.compile(r'^\((\d+)\)\s*(.*)'), '通用_層級4', 3),
]

BOLD_PATTERN = re.compile(r'\*\*(.+?)\*\*')

# ==============================================================================
# 區塊資料模型
# ==============================================================================

class Block:
    def __init__(self, style, text, ilvl=None, needs_num=False, raw_text=None):
        self.style = style
        self.text = text
        self.ilvl = ilvl               # 0~3
        self.needs_num = needs_num
        self.raw_text = raw_text or text

# ==============================================================================
# 核心排版邏輯
# ==============================================================================

def find_and_remove_anchor(doc):
    """反查錨點樣式對應的編號 ID。"""
    anchor_para = None
    for p in doc.paragraphs:
        if p.style.name == ANCHOR_STYLE:
            anchor_para = p
            break

    if anchor_para is None:
        raise RuntimeError(f"模板中找不到 '{ANCHOR_STYLE}' 樣式的錨點段落")

    # 取得 numId 與 abstractNumId
    pPr = anchor_para._element.get_or_add_pPr()
    num_id = None
    if pPr.numPr is not None and pPr.numPr.numId is not None:
        num_id = pPr.numPr.numId.val
    else:
        # 從樣式鏈反查
        num_id = _trace_num_id_from_style(doc, anchor_para.style)

    abstract_num_id = _get_abstract_num_id(doc, num_id)
    anchor_para._element.getparent().remove(anchor_para._element)
    return num_id, abstract_num_id

def _trace_num_id_from_style(doc, style):
    curr = style
    while curr:
        pPr = curr._element.find('.//w:pPr', NSMAP)
        if pPr is not None:
            numPr = pPr.find('w:numPr', NSMAP)
            if numPr is not None:
                nid = numPr.find('w:numId', NSMAP)
                if nid is not None: return int(nid.get(qn('w:val')))
        if curr.base_style: curr = curr.base_style
        else: break
    raise RuntimeError("無法從樣式鏈取得編號 ID")

def _get_abstract_num_id(doc, num_id):
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    for num in numbering.findall('.//w:num', NSMAP):
        if int(num.get(qn('w:numId'))) == num_id:
            return int(num.find('w:abstractNumId', NSMAP).get(qn('w:val')))
    raise RuntimeError(f"找不到對應的 abstractNumId (numId={num_id})")

def create_override_num(doc, abstract_num_id):
    """建立新的編號實例以重新起算。"""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    max_id = max([int(n.get(qn('w:numId'))) for n in numbering.findall('.//w:num', NSMAP)] + [0])
    new_id = max_id + 1
    
    num_el = etree.SubElement(numbering, qn('w:num'))
    num_el.set(qn('w:numId'), str(new_id))
    etree.SubElement(num_el, qn('w:abstractNumId')).set(qn('w:val'), str(abstract_num_id))
    
    for lvl in range(4):
        ov = etree.SubElement(num_el, qn('w:lvlOverride'))
        ov.set(qn('w:ilvl'), str(lvl))
        etree.SubElement(ov, qn('w:startOverride')).set(qn('w:val'), '1')
    return new_id

def setup_page_rules(doc):
    """設定行編號與懸尾控制。"""
    for sec in doc.sections:
        ln = sec._sectPr.get_or_add_lnNumType()
        ln.set(qn('w:countBy'), '1')
        ln.set(qn('w:restart'), 'newPage')

def disable_hanging_punctuation(p):
    """關閉段落懸尾。"""
    pPr = p._element.get_or_add_pPr()
    for el in pPr.findall(qn('w:overflowPunct')): pPr.remove(el)
    etree.SubElement(pPr, qn('w:overflowPunct')).set(qn('w:val'), '0')

def apply_body_indent(p, ilvl):
    """套用論述段落縮排。"""
    indent = BODY_INDENT_MAP.get(ilvl, 0)
    if indent > 0:
        p.paragraph_format.left_indent = Twips(indent)

# ==============================================================================
# 解析與寫入
# ==============================================================================

def parse_markdown(path):
    with open(path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    blocks = []
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Heading 1 & 2
        if line.startswith('# '):
            blocks.append(Block('Heading 1', line[2:].strip()))
            continue
        if line.startswith('## '):
            blocks.append(Block('Heading 2', line[3:].strip()))
            continue
            
        # 大綱前綴
        matched = False
        for pattern, style, ilvl in LEVEL_PATTERNS:
            m = pattern.match(line)
            if m:
                blocks.append(Block(style, m.group(2).strip(), ilvl=ilvl, needs_num=True))
                matched = True
                break
        if matched: continue
        
        # 預設 Normal
        blocks.append(Block('Normal', line))
    return blocks

def write_block(doc, block, num_id):
    # 處理粗體 runs
    p = doc.add_paragraph(style=block.style if block.style in [s.name for s in doc.styles] else 'Normal')
    text = block.text
    parts = []
    last = 0
    for m in BOLD_PATTERN.finditer(text):
        if m.start() > last: parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text): parts.append((text[last:], False))
    
    for t_part, is_bold in (parts if parts else [(text, False)]):
        run = p.add_run(t_part)
        if is_bold: run.bold = True
        
    # 設定編號
    if block.needs_num and block.ilvl is not None:
        pPr = p._element.get_or_add_pPr()
        for old in pPr.findall(qn('w:numPr')): pPr.remove(old)
        numPr = etree.SubElement(pPr, qn('w:numPr'))
        etree.SubElement(numPr, qn('w:ilvl')).set(qn('w:val'), str(block.ilvl))
        etree.SubElement(numPr, qn('w:numId')).set(qn('w:val'), str(num_id))
    
    disable_hanging_punctuation(p)
    return p

# ==============================================================================
# 主程式
# ==============================================================================

def main():
    parser = argparse.ArgumentParser(description='通用多層次大綱排版引擎 (v3)')
    parser.add_argument('draft', help='Markdown 路徑')
    parser.add_argument('--template', help='模板路徑')
    parser.add_argument('--output', help='輸出路徑')
    args = parser.parse_args()

    tpl_path = args.template or os.path.join(os.path.dirname(__file__), '..', 'assets', 'outline-base.docx')
    doc = Document(tpl_path)
    
    # 1. 初始化
    anchor_num_id, abstract_num_id = find_and_remove_anchor(doc)
    # 清空 body (除了 sectPr)
    body = doc.element.body
    for c in list(body):
        if c.tag != qn('w:sectPr'): body.remove(c)
    setup_page_rules(doc)
    
    # 2. 解析
    blocks = parse_markdown(args.draft)
    
    # 3. 寫入
    curr_num_id = anchor_num_id
    last_ilvl = None
    
    for b in blocks:
        # 當出現新的第一層，重設編號
        if b.style == '通用_層級1' and b.ilvl == 0:
            curr_num_id = create_override_num(doc, abstract_num_id)
            
        p = write_block(doc, b, curr_num_id)
        
        # 論述縮排
        if b.style == 'Normal' and last_ilvl is not None:
            apply_body_indent(p, last_ilvl)
            
        if b.needs_num: last_ilvl = b.ilvl
        elif b.style != 'Normal': last_ilvl = None
        
    out = args.output or (os.path.splitext(args.draft)[0] + "_output.docx")
    doc.save(out)
    print(f"成功產出: {out}")

if __name__ == "__main__":
    main()
