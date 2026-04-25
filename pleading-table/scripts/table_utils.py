#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
table_utils.py
爭點整理表格 + 聲請調查證據表建立工具。

核心策略：
1. 以 copy.deepcopy(blueprint_tbl_xml) 深拷貝傳入的藍圖表格底層 XML
2. 儲存格內段落全部刪除，再以 cell.add_paragraph() 全新建立
3. 以 style_id 強制覆寫段落樣式
4. 資料列不設 cantSplit，允許跨頁分割；需顯式移除從表頭繼承的 cantSplit / tblHeader
5. 爭點整理表三個內容欄使用純文字清單：• / 1. / 2. ...，不使用 Word 清單功能
"""

import copy
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Row


# ─────────────────────────────────────────────
# 樣式工具
# ─────────────────────────────────────────────

def _get_style_id(doc: Document, style_name: str) -> str:
    """取得 styleId；支援可能的 Mojibake / fallback style id。"""
    try:
        return doc.styles[style_name].style_id
    except KeyError:
        mangled_map = {
            "書狀_預設": ["a6", "Style1", "Ѫ_w]"],
            "爭點表_標題": ["a12", "Style15", "I_D"],
            "爭點表_清單": ["a14", "Style17", "I_M"],
            "爭點表_內容": ["a22", "Style28", "I_"],
        }
        lookups = mangled_map.get(style_name, [])
        for s in doc.styles:
            if s.style_id in lookups or s.name in lookups:
                return s.style_id
        return style_name


def _force_para_style(para, doc: Document, style_name: str) -> None:
    """直接覆寫段落樣式，必須使用 style_id。"""
    style_id = _get_style_id(doc, style_name)
    p_pr = para._element.get_or_add_pPr()
    old = p_pr.find(qn("w:pStyle"))
    if old is not None:
        p_pr.remove(old)
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), style_id)
    p_pr.insert(0, p_style)


def _force_run_font_size(para, size_pt: float) -> None:
    """覆寫段落內全部 runs 的字級。"""
    for run in para.runs:
        run.font.size = Pt(size_pt)


# ─────────────────────────────────────────────
# 儲存格工具
# ─────────────────────────────────────────────

def _delete_all_cell_paragraphs(cell) -> None:
    """刪除儲存格內全部既有段落。"""
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)


def _add_fresh_para_to_cell(cell, text: str, style_name: str, doc: Document, alignment=None):
    """在儲存格中全新建立一個段落並套用樣式。"""
    para = cell.add_paragraph()
    _force_para_style(para, doc, style_name)
    para.add_run(text)
    if alignment is not None:
        para.alignment = alignment
    return para


def _write_multiline_cell(cell, lines, style_name: str, doc: Document, alignment=None) -> None:
    """清空儲存格後，將多行文字分成多段輸出。"""
    _delete_all_cell_paragraphs(cell)
    lines = list(lines) if lines else [""]
    if not lines:
        lines = [""]
    for line in lines:
        _add_fresh_para_to_cell(cell, line, style_name, doc, alignment=alignment)


def _get_unique_cells(row) -> list:
    """去除 merged cell 重複指標後的邏輯儲存格列表。"""
    seen_tc_ids = set()
    unique_cells = []
    for c in row.cells:
        tc_id = id(c._tc)
        if tc_id not in seen_tc_ids:
            seen_tc_ids.add(tc_id)
            unique_cells.append(c)
    return unique_cells


# ─────────────────────────────────────────────
# 列屬性工具
# ─────────────────────────────────────────────

def _set_row_cant_split(row) -> None:
    """設定表格列不跨頁。"""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "1")
        tr_pr.append(cant_split)


def _set_row_tbl_header(row) -> None:
    """設定表格列為跨頁重複標題列，並同時設 cantSplit。"""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    _set_row_cant_split(row)


def _remove_row_inherit_props(row) -> None:
    """移除資料列從表頭繼承的 cantSplit / tblHeader。"""
    tr_pr = row._tr.find(qn("w:trPr"))
    if tr_pr is None:
        return
    for tag in (qn("w:cantSplit"), qn("w:tblHeader")):
        for el in tr_pr.findall(tag):
            tr_pr.remove(el)


# ─────────────────────────────────────────────
# 藍圖表格 / 列複製
# ─────────────────────────────────────────────

def clone_blueprint_table(doc: Document, blueprint_tbl_xml, anchor_element) -> Table:
    """深拷貝藍圖表格並插入到 anchor_element 之前。"""
    new_tbl = copy.deepcopy(blueprint_tbl_xml)
    anchor_element.addprevious(new_tbl)
    table = Table(new_tbl, doc)
    table.autofit = False
    return table


def _add_table_row(table: Table) -> _Row:
    """複製最後一列以新增資料列，保留格線與欄寬設定。"""
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    table._tbl.append(new_tr)
    return _Row(new_tr, table)


def _trim_table_rows(table: Table, keep_rows: int) -> None:
    """保留前 keep_rows 列，其餘刪除。"""
    while len(table.rows) > keep_rows:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)


# ─────────────────────────────────────────────
# 爭點整理表
# ─────────────────────────────────────────────

def _write_issue_list_cell(cell, items: list, doc: Document) -> None:
    """寫入原因事實 / 法律依據 / 證據或文件清單。"""
    _delete_all_cell_paragraphs(cell)

    normalized = []
    for item in items or []:
        if item is None:
            continue
        text = str(item).strip()
        if text:
            normalized.append(text)

    if not normalized:
        _add_fresh_para_to_cell(cell, "—", "爭點表_清單", doc)
        return

    if len(normalized) == 1:
        _add_fresh_para_to_cell(cell, f"•\t{normalized[0]}", "爭點表_清單", doc)
        return

    for i, text in enumerate(normalized, start=1):
        _add_fresh_para_to_cell(cell, f"{i}.\t{text}", "爭點表_清單", doc)


def _fill_issue_row(row, issue: dict, doc: Document) -> None:
    """將單一爭點內容填入資料列。"""
    _remove_row_inherit_props(row)

    cells = _get_unique_cells(row)
    if len(cells) < 4:
        raise ValueError("爭點整理表資料列欄位不足，至少需要 4 欄。")

    for cell in cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    _delete_all_cell_paragraphs(cells[0])
    _add_fresh_para_to_cell(
        cells[0],
        f"{str(issue.get('issue_number', '')).strip()}：",
        "爭點表_內容",
        doc,
    )
    _add_fresh_para_to_cell(
        cells[0],
        str(issue.get("description", "")).strip(),
        "爭點表_內容",
        doc,
    )

    _write_issue_list_cell(cells[1], issue.get("reasons", []), doc)
    _write_issue_list_cell(cells[2], issue.get("laws", []), doc)
    _write_issue_list_cell(cells[3], issue.get("evidences", []), doc)


def build_issue_table(
    doc: Document,
    issues: list,
    reason_header: str,
    anchor_para,
    blueprint_tbl_xml,
) -> Table:
    """
    建立單一爭點整理表，合併呈現事實上爭點與法律上爭點。
    欄位：
    1. 爭點
    2. {reason_header}
    3. 法律依據
    4. 證據
    """
    table = clone_blueprint_table(doc, blueprint_tbl_xml, anchor_para._element)

    # 模板應保留：表頭列 + 1 列空白資料列
    keep_rows = 2 if issues else 1
    _trim_table_rows(table, keep_rows=keep_rows)

    header_row = table.rows[0]
    _set_row_tbl_header(header_row)

    header_cells = _get_unique_cells(header_row)
    if len(header_cells) < 4:
        raise ValueError("爭點整理表藍圖欄位不足，至少需要 4 欄。")

    headers = ["爭點", reason_header, "法律依據", "證據"]
    for idx, header_text in enumerate(headers):
        cell = header_cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _write_multiline_cell(
            cell,
            [header_text],
            "爭點表_標題",
            doc,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    if not issues:
        return table

    # 先填模板內第 2 列資料列藍圖
    _fill_issue_row(table.rows[1], issues[0], doc)

    # 後續資料列再由「資料列藍圖」複製
    for issue in issues[1:]:
        row = _add_table_row(table)
        _fill_issue_row(row, issue, doc)

    return table


# ─────────────────────────────────────────────
# 聲請調查證據表
# ─────────────────────────────────────────────




def _fill_evreq_data_row(row, seq: int, item: dict, doc: Document) -> None:
    """將單筆調查證據資料填入聲請調查證據表資料列。"""
    _remove_row_inherit_props(row)

    cells = _get_unique_cells(row)
    if len(cells) < 6:
        raise ValueError("聲請調查證據表資料列欄位不足，至少需要 6 欄。")

    for cell in cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    _write_multiline_cell(
        cells[0],
        [str(seq)],
        "爭點表_內容",
        doc,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    related_issues = item.get("related_issues", []) or []
    related_lines = [str(x).strip() for x in related_issues if str(x).strip()]
    if not related_lines:
        related_lines = [""]

    _write_multiline_cell(cells[1], related_lines, "爭點表_內容", doc)
    _write_multiline_cell(
        cells[2],
        [str(item.get("investigation_item", "")).strip()],
        "爭點表_內容",
        doc,
    )
    _write_multiline_cell(
        cells[3],
        [str(item.get("target", "")).strip()],
        "爭點表_內容",
        doc,
    )

    contact = str(item.get("target_address_contact", ""))
    contact_lines = contact.split("\n") if contact else [""]
    _write_multiline_cell(cells[4], contact_lines, "爭點表_內容", doc)

    _write_multiline_cell(
        cells[5],
        [str(item.get("fact_to_prove", "")).strip()],
        "爭點表_內容",
        doc,
    )


def _build_evreq_table(
    doc: Document,
    items: list,
    blueprint_xml,
    anchor_xml_element,
) -> Table:
    """
    建立聲請調查證據表（1+N列）：
    - 第1列：欄標題列（6欄）
    - 第2列起：資料列
    """
    tbl = clone_blueprint_table(doc, blueprint_xml, anchor_xml_element)

    # 模板應保留：表頭列 + 1 列空白資料列
    keep_rows = 2 if items else 1
    _trim_table_rows(tbl, keep_rows=keep_rows)

    header_row = tbl.rows[0]
    _set_row_tbl_header(header_row)

    header_cells = _get_unique_cells(header_row)
    if len(header_cells) < 6:
        raise ValueError("聲請調查證據表藍圖欄位不足，至少需要 6 欄。")

    header_specs = [
        ["編", "號"],
        ["所涉爭點"],
        ["調查事項"],
        ["調查對象"],
        ["對象地址", "及聯絡方式"],
        ["待證事實"],
    ]

    for idx, lines in enumerate(header_specs):
        cell = header_cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _write_multiline_cell(
            cell,
            lines,
            "爭點表_標題",
            doc,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    if not items:
        return tbl

    # 先填模板內第 2 列資料列藍圖
    _fill_evreq_data_row(tbl.rows[1], 1, items[0], doc)

    # 後續資料列再由「資料列藍圖」複製
    for seq, item in enumerate(items[1:], start=2):
        row = _add_table_row(tbl)
        _fill_evreq_data_row(row, seq, item, doc)

    return tbl


def _insert_separator_paragraph_before(ref_el):
    """
    在指定 XML 元素前插入一個真正的 Word 段落，
    用來分隔兩張相鄰表格，避免 Word 自動合併。
    """
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "20")       # 1 pt
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)

    p.append(p_pr)
    ref_el.addprevious(p)
    return p


def build_evidence_request_tables(
    doc: Document,
    evidence_request: dict,
    anchor_para,
    blueprint_tbl_B_xml,
) -> Table:
    """
    建立聲請調查證據表（單一實體表格）：
    - 第 1 列：欄位標題列（6欄）
    - 第 2 列起：資料列

    最終結構：
    聲請調查證據表（欄位標題列 + 資料列）
    [anchor_para]
    """
    items = evidence_request.get("items", []) or []

    tbl = _build_evreq_table(
        doc,
        items,
        blueprint_tbl_B_xml,
        anchor_para._element,
    )

    return tbl