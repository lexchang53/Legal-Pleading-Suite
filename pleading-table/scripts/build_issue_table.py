#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_issue_table.py
爭點整理狀最終 DOCX 生成腳本。

設計原則：
- 狀首 / 狀尾移植自 draft-pleading 已驗證的邏輯
- 中文 numbering（一、/（一））從 pleading-tmpl.docx 本身反查 numId 與 abstractNumId，
  建立新的 override num，完全不依賴外部檔案
- 爭點整理表與聲請調查證據表均由 table_utils 建立
- 模板含 3 個藍圖表格：
  表格0 = 爭點整理表藍圖
  表格1 = 聲請調查證據表藍圖（僅欄位標題列 + 資料列，無標題列與提出人/日期列）
  表格2 = 高院證據清單表藍圖（序號、證據時間、證據名稱、簡要內容、待證事實、證據附卷位置（證據編號+頁碼）、備註意見）
"""

import argparse
import copy
import json
import re
import sys
from pathlib import Path

script_dir = Path(__file__).parent
# header_utils 來自 draft-pleading/scripts
draft_scripts_dir = script_dir.parent.parent / "draft-pleading" / "scripts"
if str(draft_scripts_dir) not in sys.path:
    sys.path.insert(0, str(draft_scripts_dir))

import header_utils
import table_utils

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from lxml import etree
    _HAS_LXML = True
except ImportError:
    _HAS_LXML = False


NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
ANCHOR_STYLE_OUTLINE = "通用_層級1"


# ─────────────────────────────────────────────
# 基本工具
# ─────────────────────────────────────────────

def _add_body_para(doc: Document, style_name: str, text: str):
    actual_style = header_utils._get_safe_style(doc, style_name) if hasattr(header_utils, '_get_safe_style') else style_name
    p = doc.add_paragraph(style=actual_style)
    p.add_run(text)
    return p


def enable_line_numbering(doc: Document) -> None:
    """為文件所有節啟用行編號，每頁重新起算。"""
    for section in doc.sections:
        sect_pr = section._sectPr
        ln_num_type = sect_pr.find(qn("w:lnNumType"))
        if ln_num_type is None:
            ln_num_type = OxmlElement("w:lnNumType")
            sect_pr.append(ln_num_type)
        ln_num_type.set(qn("w:countBy"), "1")
        ln_num_type.set(qn("w:restart"), "newPage")


def _add_keep_props(para) -> None:
    """在段落 pPr 中加入 keepLines + keepNext。"""
    p_pr = para._element.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


# ─────────────────────────────────────────────
# 中文 numbering 注入
# ─────────────────────────────────────────────

def _inject_chinese_numbering(doc: Document) -> int:
    """
    從已載入的 doc（來自 pleading-tmpl.docx）直接反查通用_層絆1 的 numId 與 abstractNumId，
    建立新的 w:num（帶 startOverride=1）並回傳新 numId。
    完全不依賴外部檔案。
    """
    NSMAP_LOCAL = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    ANCHOR_STYLE = "通用_層絆1"

    # --- step 1: 找到模板中套用 通用_層絆1 的段落，取得 numId ---
    num_id = None
    for p in doc.paragraphs:
        if p.style.name == ANCHOR_STYLE:
            p_pr = p._element.pPr
            if p_pr is not None and p_pr.numPr is not None and p_pr.numPr.numId is not None:
                num_id = int(p_pr.numPr.numId.val)
                break

    # fallback: 從樣式定義追查
    if num_id is None:
        for s in doc.styles:
            if s.name == ANCHOR_STYLE:
                sp_pr = s._element.find(".//w:pPr", NSMAP_LOCAL)
                if sp_pr is not None:
                    n_pr = sp_pr.find("w:numPr", NSMAP_LOCAL)
                    if n_pr is not None:
                        nid_el = n_pr.find("w:numId", NSMAP_LOCAL)
                        if nid_el is not None:
                            num_id = int(nid_el.get(qn("w:numId") or nid_el.attrib.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "1")))
                            break

    if num_id is None:
        print("[WARN] 無法從 pleading-tmpl.docx 找到 通用_層絆1 的 numId，使用 fallback numId=1")
        return 1

    # --- step 2: 反查 abstractNumId ---
    doc_npart = doc.part.numbering_part.numbering_definitions._numbering
    abstract_num_id = None
    for ne in doc_npart.findall(".//w:num", NSMAP_LOCAL):
        if int(ne.get(qn("w:numId"))) == num_id:
            anid_el = ne.find("w:abstractNumId", NSMAP_LOCAL)
            if anid_el is not None:
                abstract_num_id = int(anid_el.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))
                break

    if abstract_num_id is None:
        print("[WARN] 無法從 pleading-tmpl.docx 取得 abstractNumId，使用 fallback numId=1")
        return 1

    # --- step 3: 建立新的 w:num（startOverride=1）---
    max_num = 0
    for ne in doc_npart.findall(".//w:num", NSMAP_LOCAL):
        val = int(ne.get(qn("w:numId")))
        if val > max_num:
            max_num = val
    new_num_id = max_num + 1

    if _HAS_LXML:
        new_num = etree.SubElement(doc_npart, qn("w:num"))
        new_num.set(qn("w:numId"), str(new_num_id))
        abs_ref = etree.SubElement(new_num, qn("w:abstractNumId"))
        abs_ref.set(qn("w:val"), str(abstract_num_id))
        for ilvl_val in range(4):
            override = etree.SubElement(new_num, qn("w:lvlOverride"))
            override.set(qn("w:ilvl"), str(ilvl_val))
            so = etree.SubElement(override, qn("w:startOverride"))
            so.set(qn("w:val"), "1")
    else:
        new_num = OxmlElement("w:num")
        new_num.set(qn("w:numId"), str(new_num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_num_id))
        new_num.append(abs_ref)
        for ilvl_val in range(4):
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), str(ilvl_val))
            so = OxmlElement("w:startOverride")
            so.set(qn("w:val"), "1")
            override.append(so)
            new_num.append(override)
        doc_npart.append(new_num)

    print(f"[已完成] numbering 反查成功: numId={num_id} -> abstractNumId={abstract_num_id} -> 新 numId={new_num_id}")
    return new_num_id


def _set_num_pr(para, num_id: int, ilvl: int) -> None:
    """在段落 pPr 中設定 numPr（段落級覆寫）。"""
    p_pr = para._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)

    if _HAS_LXML:
        num_pr = etree.SubElement(p_pr, qn("w:numPr"))
        ilvl_el = etree.SubElement(num_pr, qn("w:ilvl"))
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_id_el = etree.SubElement(num_pr, qn("w:numId"))
        num_id_el.set(qn("w:val"), str(num_id))
    else:
        num_pr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl_el)
        num_pr.append(num_id_el)
        p_pr.append(num_pr)


def _add_numbered_para(doc: Document, style_name: str, text: str, num_id: int, ilvl: int):
    """新增帶 numPr 的段落，文字本身不含前綴。"""
    actual_style = header_utils._get_safe_style(doc, style_name) if hasattr(header_utils, '_get_safe_style') else style_name
    p = doc.add_paragraph(style=actual_style)
    p.add_run(text)
    if num_id is not None:
        _set_num_pr(p, num_id, ilvl)
    return p


# ─────────────────────────────────────────────
# 狀尾
# ─────────────────────────────────────────────

def _normalize_court_line(text: str) -> str:
    """正規化法院行為「法院名稱　公鑒」。"""
    m = re.match(r"^(.+?)\s*公鑒\s*$", text.strip())
    if m:
        court = m.group(1).strip()
        return f"{court}　公鑒"
    return text.strip()


def _extract_footer(docx_path: Path) -> list:
    """從既有 .docx 擷取狀尾段落（「謹狀」及其後）。"""
    src = Document(str(docx_path))
    footer_data = []
    in_footer = False

    for p in src.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        normalized = text.replace("\u3000", "").replace(" ", "")
        if normalized == "謹狀" or p.style.name == "書狀_謹狀":
            in_footer = True
        if in_footer:
            footer_data.append({"style": p.style.name, "text": text})

    print(f"[INFO] 從 '{docx_path.name}' 擷取 {len(footer_data)} 個狀尾段落")
    return footer_data


def _write_footer(doc: Document, footer_data=None) -> None:
    """
    寫入狀尾。所有狀尾段落均加 keepLines + keepNext。
    - 謹狀行：套用 書狀_謹狀
    - 法院行：套用 書狀_預設
    - 其餘：套用 書狀_簽章
    """
    import datetime

    today = datetime.date.today()
    roc_year = today.year - 1911

    def _add_footer_para(style_name: str, text: str):
        p = _add_body_para(doc, style_name, text)
        _add_keep_props(p)
        return p

    if footer_data:
        for item in footer_data:
            text = item["text"]
            normalized = text.replace("\u3000", "").replace(" ", "")
            if normalized == "謹狀" or item["style"] == "書狀_謹狀":
                _add_footer_para("書狀_謹狀", "謹狀")
            elif "公鑒" in text or item["style"] in ("書狀_法院", "書狀_預設", "書狀_預設"):
                _add_footer_para("書狀_預設", _normalize_court_line(text))
            elif header_utils.DATE_PATTERN.search(normalized):
                # 依規則：狀尾簽章日期強制加上「民國」前綴，並重新計算
                target_date = datetime.date.today() + datetime.timedelta(days=2)
                roc_year = target_date.year - 1911
                new_date_str = f"民國{roc_year}年{target_date.month}月{target_date.day}日"
                _add_footer_para("書狀_簽章", new_date_str)
            else:
                _add_footer_para("書狀_簽章", text)
        print(f"[INFO] 已寫入 {len(footer_data)} 個狀尾段落（已依規則更新日期）")
    else:
        # 重新計算預設日期：今天 + 2 日
        target_date = datetime.date.today() + datetime.timedelta(days=2)
        roc_year = target_date.year - 1911
        _add_footer_para("書狀_謹狀", "謹狀")
        _add_footer_para("書狀_預設", "臺灣高等法院　公鑒")
        _add_footer_para("書狀_簽章", "具狀人：")
        _add_footer_para("書狀_簽章", "法定代理人：")
        _add_footer_para("書狀_簽章", "訴訟代理人：")
        _add_footer_para("書狀_簽章", f"民國{roc_year}年　月　日")


# ─────────────────────────────────────────────
# 模板清空 / anchor
# ─────────────────────────────────────────────

def _clear_template_body(doc: Document) -> None:
    """清空模板所有段落與表格（保留 section properties）。"""
    body = doc.element.body
    to_remove = [child for child in body if child.tag in (qn("w:p"), qn("w:tbl"))]
    for el in to_remove:
        el.getparent().remove(el)


def _add_anchor_para(doc: Document):
    return doc.add_paragraph()


# ─────────────────────────────────────────────
# 表格後補充論述
# ─────────────────────────────────────────────

def _write_post_table_markdown(doc: Document, markdown_text: str, num_id: int) -> None:
    if not markdown_text or not markdown_text.strip():
        return

    for line in markdown_text.split("\n"):
        if not line.strip():
            doc.add_paragraph(style="書狀_預設")
            continue

        stripped = line.strip()

        if re.match(r"^[一二三四五六七八九十]+、", stripped):
            text = re.sub(r"^[一二三四五六七八九十]+、\s*", "", stripped)
            _add_numbered_para(doc, "通用_層級1", text, num_id, 0)
        elif re.match(r"^[（(][一二三四五六七八九十]+[）)]", stripped):
            text = re.sub(r"^[（(][一二三四五六七八九十]+[）)]\s*", "", stripped)
            _add_numbered_para(doc, "通用_層級2", text, num_id, 1)
        elif stripped == "謹狀":
            _add_body_para(doc, "書狀_謹狀", stripped)
        else:
            _add_body_para(doc, "書狀_預設", stripped)


# ─────────────────────────────────────────────
# payload 驗證 / 章節寫入
# ─────────────────────────────────────────────

def _get_string_list(payload: dict, key: str) -> list[str]:
    value = payload.get(key, [])
    if value is None:
        return []
    if not isinstance(value, list):
        raise ValueError(f"payload['{key}'] 必須為陣列")
    normalized = []
    for idx, item in enumerate(value):
        if not isinstance(item, str):
            raise ValueError(f"payload['{key}'][{idx}] 必須為字串")
        text = item.strip()
        if text:
            normalized.append(text)
    return normalized


def _validate_dual_mode_text(payload: dict, text_key: str, items_key: str, section_name: str) -> None:
    text = payload.get(text_key, "")
    items = _get_string_list(payload, items_key)

    if text is None:
        text = ""
    if not isinstance(text, str):
        raise ValueError(f"payload['{text_key}'] 必須為字串")

    has_text = bool(text.strip())
    has_items = bool(items)

    if has_text and has_items:
        raise ValueError(f"{section_name} 只能擇一使用 {text_key} 或 {items_key}，不得同時有內容")
    if not has_text and not has_items:
        raise ValueError(f"{section_name} 必須提供 {text_key} 或 {items_key} 其中之一")


def _validate_payload(payload: dict) -> None:
    required_root_keys = ["party_status", "reason_header", "factual_issues", "legal_issues"]
    for key in required_root_keys:
        if key not in payload:
            raise ValueError(f"payload 缺少必要欄位：{key}")

    _validate_dual_mode_text(payload, "statement_text", "statement_items", "聲明")
    _validate_dual_mode_text(payload, "undisputed_text", "undisputed_items", "不爭執事項")

    if payload["reason_header"] not in ("主張原因事實", "答辯原因事實"):
        raise ValueError("payload['reason_header'] 必須是 '主張原因事實' 或 '答辯原因事實'")

    for issues_key in ("factual_issues", "legal_issues"):
        issues = payload.get(issues_key)
        if not isinstance(issues, list):
            raise ValueError(f"payload['{issues_key}'] 必須為陣列")
        for idx, issue in enumerate(issues):
            if not isinstance(issue, dict):
                raise ValueError(f"payload['{issues_key}'][{idx}] 必須為物件")
            for field in ("issue_number", "description", "reasons", "laws", "evidences"):
                if field not in issue:
                    raise ValueError(f"payload['{issues_key}'][{idx}] 缺少必要欄位：{field}")

    er = payload.get("evidence_request")
    if er is not None:
        if not isinstance(er, dict):
            raise ValueError("payload['evidence_request'] 必須為物件或 null")
        items = er.get("items", [])
        if not isinstance(items, list):
            raise ValueError("payload['evidence_request']['items'] 必須為陣列")
        for idx, item in enumerate(items):
            if not isinstance(item, dict):
                raise ValueError(f"payload['evidence_request']['items'][{idx}] 必須為物件")
            for field in ("related_issues", "investigation_item", "target", "target_address_contact", "fact_to_prove"):
                if field not in item:
                    raise ValueError(f"payload['evidence_request']['items'][{idx}] 缺少必要欄位：{field}")


def _write_statement_section(doc: Document, payload: dict, num_id: int) -> None:
    statement_text = payload.get("statement_text", "").strip()
    statement_items = _get_string_list(payload, "statement_items")

    if statement_items:
        _add_numbered_para(doc, "通用_層級1", "聲明：", num_id, 0)
        for item in statement_items:
            _add_numbered_para(doc, "通用_層級2", item, num_id, 1)
    else:
        _add_numbered_para(doc, "通用_層級1", f"聲明：{statement_text}", num_id, 0)


def _write_undisputed_section(doc: Document, payload: dict, num_id: int) -> None:
    undisputed_text = payload.get("undisputed_text", "").strip()
    undisputed_items = _get_string_list(payload, "undisputed_items")

    if undisputed_items:
        _add_numbered_para(doc, "通用_層級1", "不爭執事項：", num_id, 0)
        for item in undisputed_items:
            _add_numbered_para(doc, "通用_層級2", item, num_id, 1)
    else:
        _add_numbered_para(doc, "通用_層級1", f"不爭執事項：{undisputed_text}", num_id, 0)


# 高院證據清單表
# ─────────────────────────────────────────────

def _set_cell_text(cell, text: str, style_name: str, doc: Document) -> None:
    """清空儲存格並以指定段落樣式填入文字。"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            pass
        p.runs[0].text = text if cell.paragraphs[0].runs else ""
        if not p.runs:
            p.add_run(text)
    else:
        np = cell.add_paragraph(text)
        try:
            np.style = doc.styles[style_name]
        except KeyError:
            pass


def _set_tbl_header(row) -> None:
    """對 table row 設定 tblHeader，使跨頁時標題列重複。"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    if trPr.find(qn("w:tblHeader")) is None:
        hdr = OxmlElement("w:tblHeader")
        trPr.append(hdr)


def build_evidence_list_table(
    doc: Document,
    evidence_list: dict,
    anchor_para,
    blueprint_tbl_xml,
) -> None:
    """
    依 evidence_list payload 產出高院證據清單表。
    """
    items = evidence_list.get("items", [])
    if not items:
        print("[INFO] evidence_list.items 為空，跳過高院證據清單表")
        return

    # 依 evidence_date 排序：0000000 排最後
    def _sort_key(item):
        d = item.get("evidence_date", "0000000")
        if d == "0000000":
            return "9999999"  # 排在最後
        return d

    sorted_items = sorted(items, key=_sort_key)

    # 複製藍圖表格（含標題列）
    new_tbl = copy.deepcopy(blueprint_tbl_xml)

    # 設定標題列（列0、列1）為跨頁重複
    rows = new_tbl.findall(qn("w:tr"))
    header_row_count = min(2, len(rows))
    for i in range(header_row_count):
        tr = rows[i]
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        if trPr.find(qn("w:tblHeader")) is None:
            hdr = OxmlElement("w:tblHeader")
            trPr.append(hdr)
        
        # 強制標題列雙向置中
        for tc in tr.findall(qn("w:tc")):
            # 垂直置中
            tcPr = tc.get_or_add_tcPr()
            vAlign = tcPr.find(qn("w:vAlign"))
            if vAlign is None:
                vAlign = OxmlElement("w:vAlign")
                tcPr.append(vAlign)
            vAlign.set(qn("w:val"), "center")
            
            # 水平置中 (對段落設定)
            for p_el in tc.findall(qn("w:p")):
                pPr = p_el.get_or_add_pPr()
                jc = pPr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    pPr.append(jc)
                jc.set(qn("w:val"), "center")

    # 取得藍圖資料列（列2）
    if len(rows) < 3:
        print("[WARN] 高院證據清單藍圖表格列數不足，跳過")
        return
    blueprint_data_row = rows[2]

    # 移除藍圖資料列（稍後逐筆複製）
    new_tbl.remove(blueprint_data_row)

    DATA_STYLE = "爭點表_內容"
    
    def _normalize_page_range(text: str) -> str:
        import re
        def _replace(match):
            p1_s, p2_s = match.group(1), match.group(2)
            try:
                if int(p2_s) > int(p1_s):
                    return f"{p1_s}~{p2_s}"
            except: pass
            return f"{p1_s}-{p2_s}"
        return re.sub(r"(\d+)\s*-\s*(\d+)", _replace, text)

    def _fill_cell(tc, text: str, align: str = "left"):
        """清空儲存格並填入文字與樣式。"""
        # 清空現有段落文字
        for p_el in tc.findall(qn("w:p")):
            for r_el in p_el.findall(qn("w:r")):
                for t_el in r_el.findall(qn("w:t")):
                    t_el.text = ""
        # 垂直置中
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)

        paras = tc.findall(qn("w:p"))
        if paras:
            p_el = paras[0]
            pPr = p_el.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p_el.insert(0, pPr)
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is None:
                pStyle = OxmlElement("w:pStyle")
                pPr.insert(0, pStyle)
            pStyle.set(qn("w:val"), DATA_STYLE)

            # 水平對齊
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center" if align == "center" else "left")
            pPr.append(jc)
            # 設文字
            runs = p_el.findall(".//" + qn("w:r"))
            if runs:
                t_els = runs[0].findall(qn("w:t"))
                if t_els:
                    t_els[0].text = text
                    if " " in text or text != text.strip():
                        t_els[0].set(
                            "{http://www.w3.org/XML/1998/namespace}space",
                            "preserve",
                        )
                else:
                    t_el = OxmlElement("w:t")
                    t_el.text = text
                    runs[0].append(t_el)
            else:
                r_el = OxmlElement("w:r")
                t_el = OxmlElement("w:t")
                t_el.text = text
                r_el.append(t_el)
                p_el.append(r_el)
        else:
            np = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            pStyle = OxmlElement("w:pStyle")
            pStyle.set(qn("w:val"), DATA_STYLE)
            pPr.append(pStyle)
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center" if align == "center" else "left")
            pPr.append(jc)
            np.append(pPr)
            nr = OxmlElement("w:r")
            nt = OxmlElement("w:t")
            nt.text = text
            nr.append(nt)
            np.append(nr)
            tc.append(np)

    for seq_idx, item in enumerate(sorted_items, start=1):
        new_row = copy.deepcopy(blueprint_data_row)
        tcs = new_row.findall(qn("w:tc"))

        # TC0: 序號 (雙向置中)
        if len(tcs) > 0:
            _fill_cell(tcs[0], str(seq_idx), align="center")
        # TC1: 證據時間 (雙向置中)
        if len(tcs) > 1:
            _fill_cell(tcs[1], item.get("evidence_date", ""), align="center")
        # TC2: 證據名稱 (垂直置中 + 靠左)
        if len(tcs) > 2:
            _fill_cell(tcs[2], item.get("evidence_name", ""), align="left")
        # TC3: 證據簡要內容 (垂直置中 + 靠左)
        if len(tcs) > 3:
            _fill_cell(tcs[3], item.get("evidence_summary", ""), align="left")
        # TC4: 待證事實 (垂直置中 + 靠左)
        if len(tcs) > 4:
            _fill_cell(tcs[4], item.get("fact_to_prove", ""), align="left")
        # TC5: 證據編號 (雙向置中)
        if len(tcs) > 5:
            _fill_cell(tcs[5], item.get("evidence_code", ""), align="center")
        # TC6: 法院卷宗頁碼 (垂直置中 + 靠左 + 範圍偵測)
        if len(tcs) > 6:
            raw_page = str(item.get("court_page", "")).strip()
            normalized_page = _normalize_page_range(raw_page)
            _fill_cell(tcs[6], normalized_page, align="left")
        # TC7: 備註意見 (垂直置中 + 靠左)
        if len(tcs) > 7:
            _fill_cell(tcs[7], item.get("remarks", ""), align="left")

        new_tbl.append(new_row)

    # 插入表格到文件 body，位置在錨點段落之前
    anchor_el = anchor_para._element
    anchor_el.addprevious(new_tbl)

    print(f"[INFO] 高院證據清單表 已生成（{len(sorted_items)} 筆）")


# ─────────────────────────────────────────────
# 自動掃描最新書狀
# ─────────────────────────────────────────────

def _find_latest_docx(directory: Path, exclude: Path = None) -> Path | None:
    def _scan(folder: Path) -> Path | None:
        candidates = [
            f for f in folder.glob("*.docx")
            if not f.name.startswith("~$")
            and not (exclude and f.resolve() == exclude.resolve())
        ]
        return max(candidates, key=lambda f: f.stat().st_mtime) if candidates else None

    result = _scan(directory)
    if result:
        return result

    data_dir = directory / "data"
    if data_dir.is_dir():
        return _scan(data_dir)

    return None


# ─────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────

def build_issue_table(
    payload: dict,
    template_path: Path,
    output_path: Path,
    header_source_path: Path = None,
) -> None:
    """生成爭點整理狀 DOCX。"""
    _validate_payload(payload)

    doc = Document(str(template_path))

    if len(doc.tables) < 2:
        raise ValueError(f"模板中表格數量不足（實際為 {len(doc.tables)} 個，至少需要 2 個藍圖表格）。")

    blueprint_issue_tbl_xml = doc.tables[0]._tbl
    blueprint_evreq_tbl_B_xml = doc.tables[1]._tbl
    blueprint_evlist_tbl_xml = doc.tables[2]._tbl if len(doc.tables) >= 3 else None

    num_id = _inject_chinese_numbering(doc)
    _clear_template_body(doc)

    if header_source_path and header_source_path.exists():
        header_data = header_utils.extract_header(header_source_path)
        header_utils.merge_and_write_header(
            doc,
            header_data=header_data,
            md_headers=None,
            is_issue_table=True,
            party_status=payload["party_status"],
        )
    else:
        header_utils.merge_and_write_header(
            doc,
            header_data=None,
            md_headers=None,
            is_issue_table=True,
            party_status=payload["party_status"],
        )

    _write_statement_section(doc, payload, num_id)
    _write_undisputed_section(doc, payload, num_id)

    _add_numbered_para(
        doc,
        "通用_層級1",
        f"{payload['party_status']}爭點整理表",
        num_id,
        0,
    )

    if "issues" in payload and payload["issues"]:
        all_issues = payload.get("issues", [])
    else:
        factual = payload.get("factual_issues", [])
        legal = payload.get("legal_issues", [])
        all_issues = factual + legal

    if all_issues:
        anchor = _add_anchor_para(doc)
        table_utils.build_issue_table(
            doc,
            all_issues,
            payload["reason_header"],
            anchor,
            blueprint_issue_tbl_xml,
        )
        anchor._element.getparent().remove(anchor._element)
    else:
        _add_numbered_para(doc, "通用_層級2", "無爭點", num_id, 1)

    er = payload.get("evidence_request")
    if er and isinstance(er, dict) and er.get("items"):
        _add_numbered_para(doc, "通用_層級1", "聲請調查證據表", num_id, 0)

        er_anchor = _add_anchor_para(doc)
        table_utils.build_evidence_request_tables(
            doc,
            er,
            er_anchor,
            blueprint_evreq_tbl_B_xml,
        )
        er_anchor._element.getparent().remove(er_anchor._element)
        print(f"[INFO] 聲請調查證據表 已生成（{len(er['items'])} 筆）")
    else:
        print("[INFO] evidence_request 不存在或 items 為空，跳過聲請調查證據表")

    post_md = payload.get("post_table_markdown", "")
    if post_md and post_md.strip():
        _write_post_table_markdown(doc, post_md, num_id)

    # 高院證據清單表：產出獨立 DOCX，不屬於爭點整理狀正文
    el_payload = payload.get("evidence_list")
    if el_payload and isinstance(el_payload, dict) and el_payload.get("items"):
        if blueprint_evlist_tbl_xml is None:
            print("[WARN] 模板中找不到第三張藍圖表格（高院證據清單），跳過")
        else:
            title_text = el_payload.get("title", "證據清單表")
            el_output_path = output_path.parent / f"{title_text}.docx"

            el_doc = Document(str(template_path))
            el_blueprint_xml = el_doc.tables[2]._tbl if len(el_doc.tables) >= 3 else None

            if el_blueprint_xml is None:
                print("[WARN] 無法從模板讀取高院證據清單藍圖，跳過")
            else:
                _clear_template_body(el_doc)

                # 寫入標題
                title_para = el_doc.add_paragraph(style="書狀_標題")
                title_para.add_run(title_text)

                # 錨點
                el_anchor = el_doc.add_paragraph()

                # 產出表格
                build_evidence_list_table(
                    el_doc,
                    el_payload,
                    el_anchor,
                    el_blueprint_xml,
                )

                # 移除錨點
                el_anchor._element.getparent().remove(el_anchor._element)

                el_doc.save(str(el_output_path))
                print(f"[輸出] 高院證據清單表儲存至：{el_output_path}")
    else:
        print("[INFO] evidence_list 不存在或 items 為空，跳過高院證據清單表")

    if header_source_path and header_source_path.exists():
        footer_data = _extract_footer(header_source_path)
        _write_footer(doc, footer_data if footer_data else None)
    else:
        _write_footer(doc, None)

    enable_line_numbering(doc)
    doc.save(str(output_path))
    print(f"[輸出] 儲存至：{output_path}")


# ─────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="以 payload JSON + pleading-tmpl.docx 生成爭點整理狀 DOCX。"
    )
    parser.add_argument("payload", help="payload JSON 路徑")
    parser.add_argument("--template", "-t", default=None)
    parser.add_argument("--output", "-o", default="爭點整理狀.docx")
    parser.add_argument("--header-source", default=None)
    args = parser.parse_args()

    payload_path = Path(args.payload)
    if not payload_path.exists():
        print(f"[錯誤] 找不到 payload：{payload_path}", file=sys.stderr)
        sys.exit(1)

    if args.template:
        template_path = Path(args.template)
    else:
        # 預設使用 draft-pleading 的共用模板（已合併爭點表藍圖與樣式）
        template_path = script_dir.parent.parent / "draft-pleading" / "assets" / "pleading-tmpl.docx"

    if not template_path.exists():
        print(f"[錯誤] 找不到模板：{template_path}", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if args.header_source:
        header_source = Path(args.header_source)
        if not header_source.exists():
            print(f"[警告] 指定書狀不存在：{header_source}，改用預設骨架")
            header_source = None
    else:
        header_source = _find_latest_docx(output_path.parent, exclude=output_path)
        if header_source:
            print(f"[自動掃描] 狀首/狀尾來源：{header_source.name}")
        else:
            print("[警告] 找不到既有書狀 .docx，使用預設骨架。")
            print("[提示] 可用 --header-source 指定書狀路徑。")

    try:
        payload = json.loads(payload_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"[錯誤] payload JSON 解析失敗：{e}", file=sys.stderr)
        sys.exit(1)

    print("[開始] 生成爭點整理狀...")
    if "issues" in payload:
        print(f" 所有爭點：{len(payload.get('issues', []))} 個")
    else:
        print(f" 事實上爭點：{len(payload.get('factual_issues', []))} 個")
        print(f" 法律上爭點：{len(payload.get('legal_issues', []))} 個")

    try:
        build_issue_table(
            payload=payload,
            template_path=template_path,
            output_path=output_path,
            header_source_path=header_source,
        )
    except Exception as e:
        import traceback
        print(f"[錯誤] 生成失敗：{e}", file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)

    print(f"[完成] 已儲存至：{output_path}")


if __name__ == "__main__":
    main()